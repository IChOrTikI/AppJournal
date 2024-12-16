import sys
import os
from PyQt5 import QtWidgets
from PyQt5.QtWidgets import QMessageBox, QListWidgetItem, QVBoxLayout, QWidget, QLabel, QPushButton, QHBoxLayout
from PyQt5 import QtGui
from PyQt5.QtCore import QThread

from docx import Document

from MainWindow import Ui_MainWindow
from EditAuthorWindow import Ui_EditAuthorWindow
from EditArticleWindow import Ui_EditArticleWindow
from EditAuthorWidget import Ui_Form  
from EditArticleWidget import Ui_Form_Article

import mysql.connector
from mysql.connector import connect, Error

class MainWindow(QtWidgets.QMainWindow, Ui_MainWindow):
    def __init__(self):

        super(MainWindow, self).__init__()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)

        #Создание БД
        self.create_db()

        # self.test_connection_create_db()

        # Первым делом открываем страницу с авторами
        self.open_page_of_authors()

        # Кнопка выхода из программы
        self.ui.button_exit.clicked.connect(self.close_app)

        # Кнопки для переключения между окнами
        self.ui.button_open_page_author.clicked.connect(self.open_page_of_authors)
        self.ui.button_open_page_articles.clicked.connect(self.open_page_of_articles)
        self.ui.button_open_page_journals.clicked.connect(self.open_page_of_journals)

        # Вызов метода для добавления элементов в ComboBox на страницк авторов
        self.add_to_combo_box_authors()    

        # Вызов метода для добавление элементов в ComboBox на странице статей    
        self.add_to_combo_box_articles()

        # Вызов метода для добавления элементов в ComboBox на странице журналов
        self.add_to_combo_box_journals()

        # Подключение кнопки к функции добавления авторов в БД
        self.ui.button_add_author.clicked.connect(self.add_author_to_db)

        # Подключение кнопки для поиска авторов по параметрам
        self.ui.button_search_authors_with_param.clicked.connect(self.search_authors_with_param)

        # Подключение кнопки к функции добавления статьи в БД
        self.ui.button_add_articles.clicked.connect(self.add_article_to_db)

        # Подключение кнопки поиска по параметрам
        self.ui.button_search_articles_with_param.clicked.connect(self.search_articles_with_param)

        # Подключение кнопки для доавбления журнала в БД
        self.ui.push_button_add_journals.clicked.connect(self.add_journal_to_db)
        
        # Поуключение кнопки для поиска журналов по параметрам
        self.ui.button_search_journals_with_param.clicked.connect(self.search_journals_with_param)

        # Подключаем метод кодга меняются параметры для посика журнала
        self.ui.combo_box_journals.currentIndexChanged.connect(self.set_mask_to_line_edit_value_param_journals)


    #  Метод для замены маски ввода для парметров поиска журнала
    def set_mask_to_line_edit_value_param_journals(self):
        if self.ui.combo_box_journals.currentText() == "Дата":
            self.ui.line_edit_value_param_journals.setInputMask('0000-00-00;_')
        else:
            self.ui.line_edit_value_param_journals.setInputMask('')
            
    # Метод для закрытия приложения
    def close_app(self):
        self.close()

    # Методы для переключения окон
    def open_page_of_authors(self):
        self.ui.Main_widgets_pages.setCurrentIndex(0)
        
        # Получаем список всех авторов, так как не передаем параметры
        list_of_authors = self.get_list_of_authors()

        # Передаем список авторов для добавления авторов на окно
        self.loading_authors_from_the_database_to_page(list_of_authors)

    def open_page_of_articles(self):
        self.ui.Main_widgets_pages.setCurrentIndex(2)

        # Вызываем функцю для заргрузки авторов в поля для выбора автора сатьи при добавлении статьи
        self.load_authors_to_combo_box_article()

        # Установление маски для ввода даты создания статьи
        self.ui.line_edit_data_article.setInputMask('0000-00-00;_')
    
        # Получение всех стаетй
        list_articles = self.get_all_articles_from_db()

        # Загрузка статей в widget с статьями
        self.load_articles_from_db_to_page_articles(list_articles)

    def open_page_of_journals(self):
        self.ui.Main_widgets_pages.setCurrentIndex(1)

        # Установление маски для ввода даты создания статьи
        self.ui.line_edit_journals_data.setInputMask('0000-00-00;_')

        # Получаем список статей который не испольузются
        list_of_articles_no_use = self.loading_all_articles_with_not_use()

        # Добавляем статьи в list widget no use
        self.insert_articles_to_widget_no_use(list_of_articles_no_use)

        # Получение всех журналов
        all_journals = self.get_all_journals_from_db()

        # Загрузка журналовв list_widget
        self.loading_journasl_to_list_widget_page_journals(all_journals)

    # Метод для добавления элементов в ComboBox на странице автор
    def add_to_combo_box_authors(self):

        self.ui.combo_box_authors.addItem("Имя")
        self.ui.combo_box_authors.addItem("Фамилия")
        self.ui.combo_box_authors.addItem("Отчество")
        self.ui.combo_box_authors.addItem("Дополнительная информация")
        self.ui.combo_box_authors.addItem("Сброс параметров")

    # Метод для добавления в ComboBox на странице статей
    def add_to_combo_box_articles(self):
        
        # self.ui.combo_box_articles.addItem("Имя автора")
        # self.ui.combo_box_articles.addItem("Фамилия автора")
        # self.ui.combo_box_articles.addItem("Отчество автора")
        self.ui.combo_box_articles.addItem("Название статьи")
        # self.ui.combo_box_articles.addItem("Дата")
        self.ui.combo_box_articles.addItem("Наука")
        self.ui.combo_box_articles.addItem("Текст статьи")
        self.ui.combo_box_articles.addItem("Сброс параметров")
        # self.ui.combo_box_articles.addItem("Текст статьи")

    # Метод для добавления в ComboBox на странице журналов
    def add_to_combo_box_journals(self):
        self.ui.combo_box_journals.addItem("Название")
        self.ui.combo_box_journals.addItem("Дата")
        self.ui.combo_box_journals.addItem("Номер")
        self.ui.combo_box_journals.addItem("Сброс параметров")

    # Метод для добавления авторов
    def add_author_to_db(self):
        first_name = self.ui.author_first_name.text()
        last_name = self.ui.author_last_name.text()
        middle_name = self.ui.author_middle_name.text()
        info_about_author = self.ui.author_additional_info.toPlainText()

        # Проверка корректности данных, Валидация данных
        errors = []

        if not first_name:
            errors.append('Необходимо указать имя автора.')
        if not last_name:
            errors.append('Необходимо указать фамилию автора.')
        if not middle_name:
            errors.append('Необходимо указать отчество автора.')
        if not info_about_author:
            errors.append('Необходимо ввести дополнительные данные об авторе.')

        if len(first_name) < 2:
            errors.append(f'Имя не может состоять из {len(first_name)} символа (символов).')
        if len(last_name) < 2:
            errors.append(f'Фамилия не может состоять из {len(last_name)} символа (символов).')
        if len(middle_name) < 5:
            errors.append(f'Отчество не может состоять из {len(middle_name)} символа (символов).')

        # Если возникла ошибка
        if errors:
            message_box = QMessageBox()
            message_box.setIcon(QMessageBox.Warning)
            message_box.setWindowTitle("Ошибка")
            message_box.setText("\n".join(errors))
            message_box.exec_()
            return
        
        # Проверка сохранения данных
        # print(first_name, last_name, middle_name, info_about_author, end='\n')

        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                cursor.execute("""INSERT INTO Authors (first_name, last_name, middle_name, info)
                VALUES (%s, %s, %s, %s); """, (first_name, last_name, middle_name, info_about_author))

                print("Пользователь добавлен")

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
                

        self.ui.author_first_name.clear()
        self.ui.author_last_name.clear()
        self.ui.author_middle_name.clear()
        self.ui.author_additional_info.clear()

        # После добавления автора обновляем список
        # Получаем список всех авторов, так как не передаем параметры
        list_of_authors = self.get_list_of_authors()

        # Передаем список авторов для добавления авторов на окно
        self.loading_authors_from_the_database_to_page(list_of_authors)

    # Метод для загрузки авторов из списка авторов на страницу с авторами (метод принимает список авторов для отображения)    
    def loading_authors_from_the_database_to_page(self, list_of_authors):
        self.ui.Widget_authors.clear()

        for author in list_of_authors:
            item = QListWidgetItem()
            item_widget = QWidget()
            name_author = QLabel(str(author[1]) + " " + str(author[2]) + " " + str(author[3]))
            separator_name = QLabel()
            text_info = QLabel(str(author[4][:95]) + "... ")
            separator_info = QLabel()

            edit_push_button = QPushButton("Редактировать")
            delete_push_button = QPushButton("Удалить")

            edit_push_button.setObjectName(str(author[0])) # Именем кнопки для редактирования будет являться id пользователя
            delete_push_button.setObjectName(str(author[0]))  # Именем кнопки для удаления будет являться id пользователя

            edit_push_button.clicked.connect(self.edit_author_push_button)
            delete_push_button.clicked.connect(self.delete_author_push_button)

            item_layout = QHBoxLayout()
            item_layout.addWidget(name_author)
            item_layout.addWidget(separator_name)
            item_layout.addWidget(text_info)
            item_layout.addWidget(separator_info)
            item_layout.addWidget(edit_push_button)
            item_layout.addWidget(delete_push_button)
            item_widget.setLayout(item_layout)

            item.setSizeHint(item_widget.sizeHint())
            self.ui.Widget_authors.addItem(item)
            self.ui.Widget_authors.setItemWidget(item, item_widget)

    # Метод для получения списка авторов из БД по параметрам (если параметры не переадются то, поиск происход по всем авторам )    
    def get_list_of_authors(self, param=None, value=None):
        connection = None
        cursor = None
        authors = []

        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                # Формируем SQL-запрос в зависимости от наличия параметров фильтрации
                
                if param == "info":
                    query = f"SELECT * FROM Authors WHERE {param} LIKE %s;"
                    cursor.execute(query, (f'%{value}%',))  # Добавляем символы подстановки
                elif param and value:
                    query = f"SELECT * FROM Authors WHERE {param} = %s;"
                    cursor.execute(query, (value,))
                else:
                    cursor.execute("SELECT * FROM Authors;")

                # Получаем все результаты при помощи fetchall
                authors = cursor.fetchall()
                print("Авторы получены.")

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()

        return authors

    # КНопки для взаимодействия с авторами на окне авторов
    # Кнопка для редактирования (функция знает id автораM открывает окно для изменения данных)
    def edit_author_push_button(self):
        sender = self.sender()
        push_button = self.findChild(QPushButton, sender.objectName())
        print(F"Кнопка редкатирования автора с ID{push_button.objectName()}")

        # Создаём новое окно MainWindow для редактирования информации о авторе
        # self.edit_window = EditAuthorWindow(str(push_button.objectName()))
        # self.edit_window.show()
        
        # Создаём окно Widget для редактирования информации о авторе
        # Передаем в окно себя и id автора
        self.edit_author_widget = EditAuthorWidget(self, str(push_button.objectName()))
        self.edit_author_widget.show()
    
    # Кнопка для удаления автора (функция знает id автора)
    def delete_author_push_button(self):

        sender = self.sender()
        push_button = self.findChild(QPushButton, sender.objectName())

        # Получаем список id статей автора
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )
            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                query = f"SELECT ID_Article FROM Authors_Articles WHERE ID_Author = %s;"
                cursor.execute(query, (push_button.objectName(),))
                id_articles = cursor.fetchall()
                print(id_articles)
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        
        # print(id_articles)

        id_for_delete = []
        
        for el in id_articles:
            id_for_delete.append(el[0])

        # print(id_for_delete)

        # Удаление статей пользователя
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )
            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                # print("id = ")
                # print(id_for_delete)
                for el in id_for_delete:
                    id_article = el # id статьи
                    # print(id_article)
                    query = f"DELETE FROM Articles WHERE ID_Article = %s;"
                    cursor.execute(query, (id_article,))

                print("Статьи удалены.")
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()



        # print(F"Кнопка удаления автора с ID{push_button.objectName()}")
        
        # Удаляем автора
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                query = f"DELETE FROM Authors WHERE ID_Author = %s;"
                cursor.execute(query, (push_button.objectName(),))

                print(f"Пользователь удалён. {push_button.objectName()}")

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()

        # После удаления автора обновляем список
        # Получаем список всех авторов, так как не передаем параметры
        list_of_authors = self.get_list_of_authors()

        # Передаем список авторов для добавления авторов на окно
        self.loading_authors_from_the_database_to_page(list_of_authors)

    # Поиск автора по параметрам
    def search_authors_with_param(self):
        
        # Получаем данные из параметров поиска
        value_text = self.ui.line_edit_author_param.text()
        value_param = self.ui.combo_box_authors.currentText()
        
        print(1)

        if value_param == "Имя":
            current_param = "first_name"
        elif value_param == "Фамилия":
            current_param = "last_name"
        elif value_param == "Отчество":
            current_param = "middle_name"
        elif value_param == "Дополнительная информация":
            current_param = "info"
        elif value_param == "Сброс параметров": # Если выбра элемент "Сброс параметров"
            value_param = None
            current_param = None
            self.ui.line_edit_author_param.clear()
            

        
        # print(value_text)
        # print(value_param)
        # value_text = str(value_text).lower()
        # print(value_text)

        info = self.get_list_of_authors(current_param, value_text)
        self.loading_authors_from_the_database_to_page(info)
    
#////////////////////

    # Метод для поиска статей по параметрам
    def search_articles_with_param(self):
        
        # Считываем значения из полей
        param = self.ui.combo_box_articles.currentText()
        value = self.ui.line_edit_articles_param.text()

        if param == "Название статьи":
            current_param = "name"
        # elif param == "Дата":
        #     current_param = "data_create"
        elif param == "Наука":
            current_param = "science"
        elif param == "Текст статьи":
            current_param = "text_article"
        elif param == "Сброс параметров": # Если выбра элемент "Сброс параметров"
            param = None
            current_param = None
            self.ui.line_edit_author_param.clear()

        # Получение списка авторов при поиске с параметрос
        list_of_articles_with_param = self.loading_articles_from_db_with_param(current_param, value)

         # Загрузка статей в widget с статьями
        self.load_articles_from_db_to_page_articles(list_of_articles_with_param)

    # Загрузка данных о статьях из БД по парметрам
    def loading_articles_from_db_with_param(self, param=None, value=None):
        connection = None
        cursor = None
        authors = []

        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                # Формируем SQL-запрос в зависимости от наличия параметров фильтрации


                if param == "text_article":
                    query = f"""
                        SELECT
                            Authors.ID_Author,
                            Authors.first_name,
                            Authors.last_name,
                            Articles.ID_Article,
                            Articles.name,
                            Articles.date_create,
                            Articles.science,
                            Articles.text_article
                        FROM
                            Articles
                        INNER JOIN
                            Authors_Articles ON Articles.ID_Article = Authors_Articles.ID_Article
                        INNER JOIN
                            Authors ON Authors_Articles.ID_Author = Authors.ID_Author
                        WHERE 
                            Articles.{param} LIKE %s;
                        """
                    cursor.execute(query, (f'%{value}%',))  # Добавляем символы подстановки
                elif param and value:
                    query = f"""
                        SELECT
                            Authors.ID_Author,
                            Authors.first_name,
                            Authors.last_name,
                            Articles.ID_Article,
                            Articles.name,
                            Articles.date_create,
                            Articles.science,
                            Articles.text_article
                        FROM
                            Articles
                        INNER JOIN
                            Authors_Articles ON Articles.ID_Article = Authors_Articles.ID_Article
                        INNER JOIN
                            Authors ON Authors_Articles.ID_Author = Authors.ID_Author
                        WHERE 
                            Articles.{param} = %s;
                        """
                    cursor.execute(query, (value,))
                else:
                    query = f"""
                        SELECT
                            Authors.ID_Author,
                            Authors.first_name,
                            Authors.last_name,
                            Articles.ID_Article,
                            Articles.name,
                            Articles.date_create,
                            Articles.science,
                            Articles.text_article
                        FROM
                            Articles
                        INNER JOIN
                            Authors_Articles ON Articles.ID_Article = Authors_Articles.ID_Article
                        INNER JOIN
                            Authors ON Authors_Articles.ID_Author = Authors.ID_Author;
                        """
                    cursor.execute(query)

                # Получаем все результаты при помощи fetchall
                authors = cursor.fetchall()
                print("Авторы получены.")

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()

        return authors

    # Метод добавление статьи в БД
    def add_article_to_db(self):

        # Получаем выбранного автора
        text = self.ui.combo_box_authors_article.currentText()

        # Из текстовой информации делаем массив с именем и фамилией
        arr_info_about_current_author = text.split()

        first_name = arr_info_about_current_author[0]
        last_name = arr_info_about_current_author[1]
        middle_name = arr_info_about_current_author[2]
        first_name = str(first_name)
        last_name = str(last_name)
        middle_name = str(middle_name)

        # Получение id автора
        id_author_for_articel = self.get_author_id_with_first_last_middle_name(first_name, last_name, middle_name)

        # Получение данных из полей ввода для статьи
        name_articel = self.ui.line_edit_name_article.text()
        data_articel = self.ui.line_edit_data_article.text() 
        science_articel = self.ui.line_edit_scince_article.text()
        text_article = self.ui.text_edit_text_article.toPlainText()

        # Валмдация данных
        # Валидация даты
        if not self.validate_date(data_articel):
            QMessageBox.warning(self, "Ошибка", "Введите корректную дату в формате ГГГГ-ММ-ДД или РЕАЛЬНУЮ дату")
            self.ui.line_edit_data_article.clear()  # Очищаем поле ввода
            return
        elif not name_articel:
            QMessageBox.warning(self, "Ошибка", "Введите навзание статьи")
            return
        elif not science_articel:
            QMessageBox.warning(self, "Ошибка", "Введите навзание науки")
            return
        elif not text_article:
            QMessageBox.warning(self, "Ошибка", "Введите текст статьи")
            return
        elif len(name_articel) >= 255:
            QMessageBox.warning(self, "Ошибка", "Введите навзание стать меньеше 255 символов")
            return
        elif len(science_articel) >= 255:
            QMessageBox.warning(self, "Ошибка", "Введите навзание науки меньше 255 символов")
            return
        

        # Добавляем статью в БД
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                cursor.execute("""INSERT INTO Articles (name, date_create, science, text_article)
                VALUES (%s, %s, %s, %s); """, (name_articel, data_articel, science_articel, text_article))

                print("Статья добавлена")

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()

        # После добавления статьи очищаем поля ввода
        self.ui.line_edit_name_article.clear()
        self.ui.line_edit_data_article.clear()
        self.ui.line_edit_scince_article.clear()
        self.ui.text_edit_text_article.clear()

        # Связываем статью с выбранным пользователем
        id_article = self.get_id_article_with_name_date_science_text(name_articel, data_articel, science_articel, text_article)
        
        self.connect_author_and_article_in_db(id_author_for_articel, id_article)

        # Получение всех стаетй
        list_articles = self.get_all_articles_from_db()

        # Загрузка статей в widget с статьями
        self.load_articles_from_db_to_page_articles(list_articles)

        # Проверка даных для статьи
        # print(id_author_for_articel)
        # print(name_articel)
        # print(data_articel)
        # print(science_articel)
        # print(text_article)

        # Проверка
        # print(arr_info_about_current_author) 
        # print(text)

    # Метод загрузки авторов в combo box с авторами на странице статей
    def load_authors_to_combo_box_article(self):

        # Убираем старые элементы 
        self.ui.combo_box_authors_article.clear()

        # Поолучаем список всех авторов из БД
        all_authors = self.get_list_of_authors()

        arr_authors = []

        for author in all_authors:
            el_author = author[1] + " " + author[2] + " " + author[3]
            arr_authors.append(el_author)
        
        # Добавляем авторов в combobox
        self.ui.combo_box_authors_article.addItems(arr_authors)        

        # print(all_authors)
        # print(arr_authors)

    # Получение автора по его ФИО
    def get_author_id_with_first_last_middle_name(self, first, last, middle):
        connection = None
        cursor = None

        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                # Формируем SQL-запрос по параметрам

                query = f"SELECT * FROM Authors WHERE first_name = %s AND last_name = %s AND middle_name = %s"
                cursor.execute(query, (first, last, middle),)

                # Получаем все результаты при помощи fetchall
                authors = cursor.fetchone()
                
                author_id = authors[0]

                print("Автор получен.")

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()

        return author_id
    
    # Получение id статьи по её данным
    def get_id_article_with_name_date_science_text(self, name, date, science, text):
        # Получаем id статьи по её данным
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                cursor.execute("""SELECT * FROM Articles WHERE name = %s AND date_create = %s AND science = %s AND text_article = %s""", (name, date, science, text))
                articles = cursor.fetchone()
                articles_id = articles[0]
                # print(articles)
                # print("Статья получена")

                return articles_id

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
    
    # Соединение автора со статьеё
    def connect_author_and_article_in_db(self, id_author, id_article):
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                cursor.execute("""INSERT INTO Authors_Articles (ID_Author,  ID_Article) 
                VALUES (%s, %s); """, (id_author, id_article))
                print("Автор и Статья подключены")

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()

    # Загрузка статей из БД на страницу статей
    def load_articles_from_db_to_page_articles(self, list_of_articles):
        
        # Вызываем метод для получения всех статей с их авторами
        # list_of_articles = self.get_all_articles_from_db()

        # Выводим список с сатьями на экран
        print(list_of_articles)

        self.ui.list_widget_articles.clear()

        for article in list_of_articles:
            item = QListWidgetItem()
            item_widget = QWidget()
            name_author = QLabel(str(article[1]) + " " + str(article[2]))
            name_article = QLabel(str(article[4]))
            date_create_article = QLabel(str(article[5]))
            science_of_article = QLabel(str(article[6]))
            text_of_article = str(article[7])
            if len(text_of_article) > 35:
                text_of_article = QLabel(str(article[7][:35]) + "...")
            else:
                text_of_article = QLabel(str(article[7]))
                

            # separator_name = QLabel()

            edit_push_button = QPushButton("Редактировать")
            delete_push_button = QPushButton("Удалить")
            view_push_button = QPushButton("Смотреть")

            edit_push_button.setObjectName(str(article[3])) # Именем кнопки для редактирования будет являться id статьи
            delete_push_button.setObjectName(str(article[3]))  # Именем кнопки для удаления будет являться id статьи
            view_push_button.setObjectName(str(article[3])) # Именем кнопки для просмотра дубет являться id статьи

            edit_push_button.clicked.connect(self.edit_article_push_button)
            delete_push_button.clicked.connect(self.delete_article_push_button)
            view_push_button.clicked.connect(self.view_article_push_button)

            item_layout = QHBoxLayout()

            item_layout.addWidget(name_author)
            item_layout.addWidget(name_article)
            item_layout.addWidget(date_create_article)
            item_layout.addWidget(science_of_article)
            item_layout.addWidget(text_of_article)
            item_layout.addWidget(edit_push_button)
            item_layout.addWidget(delete_push_button)
            item_layout.addWidget(view_push_button)
            item_widget.setLayout(item_layout)

            item.setSizeHint(item_widget.sizeHint())
            self.ui.list_widget_articles.addItem(item)
            self.ui.list_widget_articles.setItemWidget(item, item_widget)

    #  Метод для редактирования статьи
    def edit_article_push_button(self):
        sender = self.sender()
        push_button = self.findChild(QPushButton, sender.objectName())
        print("Конпка редкатирования статьи с id : ", push_button.objectName())

        # Создаём окно Widget для редактирования информации о статье
        # Передаем в окно себя и id статьи
        self.edit_article_widget = EditArticleWidget(self, str(push_button.objectName()))
        self.edit_article_widget.show()

    # Метод для удаления статьи
    def delete_article_push_button(self):
        sender = self.sender()
        push_button = self.findChild(QPushButton, sender.objectName())
        print("Кнопка удаления статьи с id : ", push_button.objectName())

        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                cursor.execute("""DELETE FROM Articles WHERE ID_Article = %s """, (push_button.objectName(),))
                print("Статья удалена")

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()

        # Получение всех стаетй
        list_articles = self.get_all_articles_from_db()

        # Загрузка статей в widget с статьями
        self.load_articles_from_db_to_page_articles(list_articles)

    # Метод для просмотра статьи
    def view_article_push_button(self):
        sender = self.sender()
        push_button = self.findChild(QPushButton, sender.objectName())
        print("Кнопка просмотра сатьи c id : ", push_button.objectName())

        onnection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                query = """
                SELECT
                    Authors.ID_Author,
                    Authors.first_name,
                    Authors.last_name,
                    Articles.ID_Article,
                    Articles.name,
                    Articles.date_create,
                    Articles.science,
                    Articles.text_article
                FROM
                    Articles
                INNER JOIN
                    Authors_Articles ON Articles.ID_Article = Authors_Articles.ID_Article
                INNER JOIN
                    Authors ON Authors_Articles.ID_Author = Authors.ID_Author
                WHERE
                    Articles.ID_Article = %s;
                """

                cursor.execute(query, (push_button.objectName(),))
                items = cursor.fetchone()
                # print(items)
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()

        # Создаем новый документ
        doc = Document()
        doc.add_heading(f"{items[4]}")
        doc.add_paragraph(f"Автор: {items[1]} {items[2]}")
        doc.add_paragraph(f"Дата: {items[5]}")
        doc.add_paragraph(f"Наука: {items[6]}")
        doc.add_paragraph(f"{items[7]}")

        # Получаем путь к директории, где находится текущий скрипт
        project_path = os.path.dirname(os.path.abspath(sys.modules['__main__'].__file__))

        # Указываем имя файла .docx
        file_name = f'article_{push_button.objectName()}.docx'  # Замените на имя вашего файла

        # Полный путь к файлу
        file_path = os.path.join(project_path, file_name)

        # Сохраняем документ (перезаписывает, если файл уже существует)
        doc.save(file_path)

        # Открываем файл в Microsoft Word
        os.startfile(file_path)

    # Получение всех статей из БД
    def get_all_articles_from_db(self):
        onnection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                cursor.execute("""
                SELECT
                    Authors.ID_Author,
                    Authors.first_name,
                    Authors.last_name,
                    Articles.ID_Article,
                    Articles.name,
                    Articles.date_create,
                    Articles.science,
                    Articles.text_article
                FROM
                    Articles
                INNER JOIN
                    Authors_Articles ON Articles.ID_Article = Authors_Articles.ID_Article
                INNER JOIN
                    Authors ON Authors_Articles.ID_Author = Authors.ID_Author;
                """)
                
                items = cursor.fetchall()
                # print(items)
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        
        return items

    # Проверка корректоности даты на странице статей 
    def validate_date(self, date_string):
        # Проверяем, что строка имеет длину 10 символов (формат ГГГГ-ММ-ДД)
        if len(date_string) != 10:
            return False

        year = date_string[:4]
        month = date_string[5:7]
        day = date_string[8:]

        # Проверка даты
        # print(year)
        # print(month)
        # print(day)

        # Проверка года
        if int(year) >= 1000:
            # print(1)
            # Проверка месяца
            if int(month) >= 1 and int(month) <= 12:
                # Проверка дня
                if int(day) >= 1 and int(day) <= 31:
                    # Дополнительная проверка на количество дней в месяце
                    if month in ['04', '06', '09', '11'] and int(day) == 31:
                        return False  # Апрель, Июнь, Сентябрь, Ноябрь имеют максимум 30 дней
                    if month == '02':
                        # Проверка на високосный год
                        if (int(year) % 4 == 0 and int(year) % 100 != 0) or (int(year) % 400 == 0):
                            if int(day) > 29:
                                return False
                        else:
                            if int(day) > 28:
                                return False
                    return True
        return False

#////////////////////////////////////////

    # Метод для добавления журнала
    def add_journal_to_db(self):

        # Получаем значения из полей
        name_journal = self.ui.line_edit_journals_name.text()
        date_journal = self.ui.line_edit_journals_data.text()
        number_journal = self.ui.line_edit_journals_number.text()

        # Валидация
        if not self.validate_date(date_journal):
            QMessageBox.warning(self, "Ошибка", "Введите корректную дату в формате ГГГГ-ММ-ДД или РЕАЛЬНУЮ дату")
            self.ui.line_edit_journals_data.clear()  # Очищаем поле ввода
            return
        elif len(name_journal) >= 255:
            QMessageBox.warning(self, "Ошибка", "Введите название журнала поменьше")
            return
        elif not name_journal:
            QMessageBox.warning(self, "Ошибка", "Введите название журнала")
            return
        elif not number_journal.isdigit():
            QMessageBox.warning(self, "Ошибка", "Номер журнала должен быть числом")
            return
        elif int(number_journal) >= 2147483647:
            QMessageBox.warning(self, "Ошибка", "Введите число поменьше")
            return

        self.add_journals_to_db(name_journal, date_journal, number_journal)

        self.ui.line_edit_journals_name.clear()
        self.ui.line_edit_journals_data.clear()
        self.ui.line_edit_journals_number.clear()


        # Получение всех журналов
        all_journals = self.get_all_journals_from_db()

        # Загрузка журналовв list_widget
        self.loading_journasl_to_list_widget_page_journals(all_journals)

        # Получени id созданного журнала
        id_journal = self.get_id_journal(name_journal, date_journal, number_journal)
        id_journal = id_journal[0]

        # Выбранные статьи
        # print("Выбранные статьи : ")
        list_name = []
        # Получаем количество элементов в QListWidget
        item_count = self.ui.list_wodget_use_articles.count()
        # Проходим по всем элементам
        for index in range(item_count):
            item = self.ui.list_wodget_use_articles.item(index)  # Получаем элемент по индексу
            if item is not None:
                article_name = item.text()  # Получаем текст элемента
                list_name.append(article_name)
                # print(article_name)  # Выводим имя статьи

        # Получаем списко id статей для текущего журнала по именам статей
        list_id_articles = self.get_id_articles(list_name)

        # print("Информация:")
        # print(id_journal)
        # print(list_id_articles)

        # Соединяем статьи с журналом
        self.connect_article_with_journal(id_journal, list_id_articles)

        # Меняем статус статей
        self.change_status_articles(list_id_articles)

        # Обновление данных

        # Получение всех журналов
        all_journals = self.get_all_journals_from_db()

        # Загрузка журналов list_widget
        self.loading_journasl_to_list_widget_page_journals(all_journals)

        # # Получаем список статей который не испольузются
        list_of_articles_no_use = self.loading_all_articles_with_not_use()

        # # Добавляем статьи в list widget no use
        self.insert_articles_to_widget_no_use(list_of_articles_no_use)
        
        self.ui.list_wodget_use_articles.clear()

        # # Получение всех журналов
        # all_journals = self.get_all_journals_from_db()

        # # Загрузка журналовв list_widget
        # self.loading_journasl_to_list_widget_page_journals(all_journals)

        # # 
        # self.list_wodget_use_articles.clear()



        # # Не выбранные статьи
        # print("Не выбранные статьи : ")
        # # Получаем количество элементов в QListWidget
        # item_count = self.ui.list_wodget_no_use_articles.count()
        # # Проходим по всем элементам
        # for index in range(item_count):
        #     item = self.ui.list_wodget_no_use_articles.item(index)  # Получаем элемент по индексу
        #     if item is not None:
        #         article_name = item.text()  # Получаем текст элемента
        #         print(article_name)  # Выводим имя статьи

    # Получаем все статьи которые ещё не исопльзуются 
    def loading_all_articles_with_not_use(self):
        onnection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                cursor.execute("""
                SELECT ID_Article, name FROM Articles WHERE isUse = 0;
                """)
                list_of_articles_no_use = cursor.fetchall()
                # print(items)
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        
        return list_of_articles_no_use

    def insert_articles_to_widget_no_use(self, list_articles):
        # Очищаем list widget no use
        self.ui.list_wodget_no_use_articles.clear()
        self.ui.list_wodget_use_articles.clear()


        # Добавляем элементы в list widget no use
        for article in list_articles:
            item = QListWidgetItem(str(article[1]))  # Устанавливаем текст элемента
            item_widget = QWidget()
            separator_name = QLabel()
            name_article = QLabel(str(article[1]))

            add_push_button = QPushButton("+")
            add_push_button.setObjectName(str(article[0]))  # Именем кнопки для переноса элемента в другой list widget use

            # Передаем текущий item в качестве аргумента в метод
            add_push_button.clicked.connect(lambda checked, item=item: self.add_article_push_button(item))

            item_layout = QHBoxLayout()
            item_layout.addWidget(name_article)
            item_layout.addWidget(separator_name)
            item_layout.addWidget(add_push_button)

            item_widget.setLayout(item_layout)

            item.setSizeHint(item_widget.sizeHint())
            self.ui.list_wodget_no_use_articles.addItem(item)
            self.ui.list_wodget_no_use_articles.setItemWidget(item, item_widget)

    def add_article_push_button(self, item):
        # Удаляем элемент из list widget no use
        row = self.ui.list_wodget_no_use_articles.row(item)
        if row != -1:
            article_name = item.text()  # Получаем текст элемента
            self.ui.list_wodget_no_use_articles.takeItem(row)

            # Создаем новый элемент для второго списка
            new_item = QListWidgetItem(article_name)  # Устанавливаем текст нового элемента
            new_item_widget = QWidget()

            remove_push_button = QPushButton("-")
            remove_push_button.setObjectName(article_name)  # Именем кнопки для удаления элемента
            separator_name = QLabel()
            remove_push_button.clicked.connect(lambda checked, item=new_item: self.remove_article_push_button(item, article_name))

            item_layout = QHBoxLayout()
            item_layout.addWidget(QLabel(article_name))  # Можно добавить название статьи
            item_layout.addWidget(separator_name)
            item_layout.addWidget(remove_push_button)

            new_item_widget.setLayout(item_layout)

            new_item.setSizeHint(new_item_widget.sizeHint())
            self.ui.list_wodget_use_articles.addItem(new_item)
            self.ui.list_wodget_use_articles.setItemWidget(new_item, new_item_widget)

    def remove_article_push_button(self, item, article_name):
        # Удаляем элемент из list widget use
        row = self.ui.list_wodget_use_articles.row(item)
        if row != -1:
            self.ui.list_wodget_use_articles.takeItem(row)

            # Создаем новый элемент для возвращения в первый список
            return_item = QListWidgetItem(article_name)  # Устанавливаем текст возвращаемого элемента
            return_widget = QWidget()

            add_push_button = QPushButton("+")
            add_push_button.setObjectName(article_name)  # Именем кнопки для переноса обратно
            separator_name = QLabel()
            add_push_button.clicked.connect(lambda checked, item=return_item: self.add_article_push_button(return_item))

            return_layout = QHBoxLayout()
            return_layout.addWidget(QLabel(article_name))  # Можно добавить название статьи
            return_layout.addWidget(separator_name)
            return_layout.addWidget(add_push_button)

            return_widget.setLayout(return_layout)

            return_item.setSizeHint(return_widget.sizeHint())
            self.ui.list_wodget_no_use_articles.addItem(return_item)
            self.ui.list_wodget_no_use_articles.setItemWidget(return_item, return_widget)

    def add_journals_to_db(self, name, data, number):
        onnection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                query = """INSERT INTO Journals (name, date_create, number_journal) VALUES (%s, %s, %s);"""

                cursor.execute(query, (name, data, number,))

                print("Журнал добавлен")

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
    
    def get_all_journals_from_db(self, param=None, value=None):
        onnection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()
                
                if param and value:
                    query = f"""SELECT * FROM Journals WHERE {param} = %s;"""
                    cursor.execute(query, (value,))
                else:
                    cursor.execute(""" SELECT * FROM Journals; """)
                
                items = cursor.fetchall()

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        
        return items

    def loading_journasl_to_list_widget_page_journals(self, list_journals):
        self.ui.list_widget_journals.clear()

        print("all_journals")
        print(list_journals)

        for journal in list_journals:
            item = QListWidgetItem()
            item_widget = QWidget()
            name_journal = QLabel(str(journal[1]))
            separator_name = QLabel()
            date_journal = QLabel(str(journal[2]))
            separator_date = QLabel()
            number_journal = QLabel(str(journal[3]))
            separator_button = QLabel()

            delete_push_button = QPushButton("Удалить")
            view_push_button = QPushButton("Смотреть")

            delete_push_button.setObjectName(str(journal[0]))  # Именем кнопки для удаления будет являться id журнала
            view_push_button.setObjectName(str(journal[0]))  # Именем кнопки для просмотра будет являться id журнала

            delete_push_button.clicked.connect(self.delete_journal_push_button)
            view_push_button.clicked.connect(self.view_journal_push_button)

            item_layout = QHBoxLayout()
            item_layout.addWidget(name_journal)
            item_layout.addWidget(separator_name)
            item_layout.addWidget(date_journal)
            item_layout.addWidget(separator_date)
            item_layout.addWidget(number_journal)
            item_layout.addWidget(separator_button)
            item_layout.addWidget(delete_push_button)
            item_layout.addWidget(view_push_button)
            item_widget.setLayout(item_layout)

            item.setSizeHint(item_widget.sizeHint())
            self.ui.list_widget_journals.addItem(item)
            self.ui.list_widget_journals.setItemWidget(item, item_widget)
    
    def connect_article_with_journal(self, id_journal, list_id_articles):

        for id_article in list_id_articles:
            onnection = None
            cursor = None
            try:
                connection = connect(
                    host="sql7.freesqldatabase.com",
                    user="sql7751998",
                    password="7kPDaYU2TX",
                    database="sql7751998" # Имя базы данных
                )

                if connection.is_connected():
                    print("Успешное подключение")
                    cursor = connection.cursor()

                    query = """INSERT INTO Articles_Journals (ID_Journal, ID_Article) VALUES (%s, %s);"""

                    cursor.execute(query, (id_journal, id_article,))

                    # print(f"Статья {id_article} соединена с {id_journal}")

            except Error as e:
                print(f"Ошибка подключения: {e}")
            finally:
                if cursor:
                    cursor.close()
                if connection:
                    connection.commit()
                    connection.close()

    def get_id_journal(self, name, date_create, number_journal):
        onnection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                query = """SELECT ID_journal FROM Journals WHERE name = %s AND date_create = %s AND  number_journal = %s;"""

                cursor.execute(query, (name, date_create, number_journal,))
                
                id_journal = cursor.fetchone()

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        
        return id_journal

    def get_id_articles(self, list_name_articles):
        
        list_id_articles = []

        for name_article in list_name_articles:
            onnection = None
            cursor = None
            try:
                connection = connect(
                    host="sql7.freesqldatabase.com",
                    user="sql7751998",
                    password="7kPDaYU2TX",
                    database="sql7751998" # Имя базы данных
                )

                if connection.is_connected():
                    print("Успешное подключение")
                    cursor = connection.cursor()

                    query = """SELECT ID_Article FROM Articles WHERE name = %s;"""

                    cursor.execute(query, (name_article,))
                
                    id_article = cursor.fetchone()
                    id_article = id_article[0] 

                    list_id_articles.append(id_article)

            except Error as e:
                print(f"Ошибка подключения: {e}")
            finally:
                if cursor:
                    cursor.close()
                if connection:
                    # connection.commit()
                    connection.close()
        
        return list_id_articles

    def change_status_articles(self, list_id_articles):
        
        for id_article in list_id_articles:
            onnection = None
            cursor = None
            try:
                connection = connect(
                    host="sql7.freesqldatabase.com",
                    user="sql7751998",
                    password="7kPDaYU2TX",
                    database="sql7751998" # Имя базы данных
                )

                if connection.is_connected():
                    print("Успешное подключение")
                    cursor = connection.cursor()

                    query = """UPDATE Articles SET isUse = 1 WHERE ID_Article = %s"""

                    cursor.execute(query, (id_article,))


            except Error as e:
                print(f"Ошибка подключения: {e}")
            finally:
                if cursor:
                    cursor.close()
                if connection:
                    connection.commit()
                    connection.close()

    def search_journals_with_param(self):
        
        # Считываем данные из полей
        param = self.ui.combo_box_journals.currentText()
        value = self.ui.line_edit_value_param_journals.text()

        if param == "Название":
            current_param = "name"
        elif param == "Дата":
            if not self.validate_date(value):
                QMessageBox.warning(self, "Ошибка", "Введите корректную дату в формате ГГГГ-ММ-ДД или РЕАЛЬНУЮ дату")
                return
            current_param = "date_create" 
        elif param == "Номер":
            current_param = "number_journal"
        elif param == "Сброс параметров":
            current_param = None
            value = None

        list_journals = self.get_all_journals_from_db(current_param, value)

        self.loading_journasl_to_list_widget_page_journals(list_journals)

    def delete_journal_push_button(self):
        sender = self.sender()
        push_button = self.findChild(QPushButton, sender.objectName())

        onnection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()


                
                query = """SELECT ID_Article FROM Articles_Journals WHERE ID_Journal = %s"""
                cursor.execute(query, (sender.objectName(),))
                id_articles = cursor.fetchall()
                print(id_articles)

                list_id_articles = []

                for el in id_articles:
                    list_id_articles.append(el[0])

                print(list_id_articles)

                # return

                query = """DELETE FROM Journals WHERE ID_journal = %s"""
                cursor.execute(query, (sender.objectName(),))

                for el in list_id_articles:
                    query = """UPDATE Articles SET isUse = 0 WHERE ID_Article = %s"""
                    cursor.execute(query, (el,))

                print("Журнал удалён.")

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        
        # # Получение всех журналов
        # all_journals = self.get_all_journals_from_db()
        # print(1)

        # # Загрузка журналовв list_widget
        # self.loading_journasl_to_list_widget_page_journals(all_journals)

        # Получение всех журналов
        all_journals = self.get_all_journals_from_db()

        # Загрузка журналов list_widget
        self.loading_journasl_to_list_widget_page_journals(all_journals)

        # # Получаем список статей который не испольузются
        list_of_articles_no_use = self.loading_all_articles_with_not_use()

        # # Добавляем статьи в list widget no use
        self.insert_articles_to_widget_no_use(list_of_articles_no_use)
        
        self.ui.list_wodget_use_articles.clear()

    def view_journal_push_button(self):
        pass

#///////////////////////////////////////
    def create_report(self):
        pass

    # Создание БД
    def create_db(self):
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                # Создание БД если её нет
                # cursor.execute(""" CREATE DATABASE IF NOT EXISTS AppJournal; """)
                
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS Authors (
                        ID_Author INT NOT NULL AUTO_INCREMENT PRIMARY KEY,
                        first_name VARCHAR(150) NOT NULL,
                        last_name VARCHAR(150) NOT NULL,
                        middle_name VARCHAR(150) NOT NULL,
                        info TEXT NOT NULL,
                        count_articles INT NOT NULL DEFAULT 0
                    );
                """)


                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS Articles (
                        ID_Article INT NOT NULL AUTO_INCREMENT PRIMARY KEY,
                        name VARCHAR(255) NOT NULL,
                        date_create DATE NOT NULL,
                        science VARCHAR(255) NOT NULL,
                        text_article TEXT NOT NULL,
                        isUse BOOLEAN NOT NULL DEFAULT FALSE
                    );
                """)

                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS Authors_Articles (
                        ID INT NOT NULL AUTO_INCREMENT PRIMARY KEY,
                        ID_Author INT,
                        ID_Article INT UNIQUE,
                        FOREIGN KEY (ID_Author) REFERENCES Authors(ID_Author) ON DELETE CASCADE,
                        FOREIGN KEY (ID_Article) REFERENCES Articles(ID_Article) ON DELETE CASCADE
                    );
                """)

                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS Journals (
                        ID_Journal INT NOT NULL AUTO_INCREMENT PRIMARY KEY,
                        name VARCHAR(255) NOT NULL,
                        date_create DATE NOT NULL,
                        number_journal INT NOT NULL
                    );
                """)

                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS Articles_Journals (
                        ID INT NOT NULL AUTO_INCREMENT PRIMARY KEY,
                        ID_Journal INT,
                        ID_Article INT UNIQUE,
                        FOREIGN KEY (ID_Journal) REFERENCES Journals(ID_Journal) ON DELETE CASCADE,
                        FOREIGN KEY (ID_Article) REFERENCES Articles(ID_Article) ON DELETE CASCADE
                    );
                """)

                print("Таблицы успешно созданы.")

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()

    # Ткстовое подключение к БД
    def test_connection_create_db(self):
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql12.freesqldatabase.com",
                user="sql12747795",
                password="ujFlrPM1xw",
                database="sql12747795"
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS authors (
                        id INT AUTO_INCREMENT PRIMARY KEY,
                        firstName VARCHAR(150) NOT NULL,
                        lastName VARCHAR(150) NOT NULL,
                        middleName VARCHAR(150),
                        info TEXT NOT NULL,
                        count_articles INT DEFAULT 0  
                )   ;
                """)

                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS articles (
                        id INT AUTO_INCREMENT PRIMARY KEY,
                        title VARCHAR(255) NOT NULL,
                        date DATE NOT NULL,
                        science VARCHAR(100) NOT NULL,
                        text TEXT NOT NULL
                    );
                """)

                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS journals (
                        id INT AUTO_INCREMENT PRIMARY KEY,
                        title VARCHAR(255) NOT NULL,
                        date DATE NOT NULL,
                        number INT NOT NULL
                    );
                """)

                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS author_article (
                        authorId INT,
                        articleId INT,
                        PRIMARY KEY (authorId, articleId),
                        FOREIGN KEY (authorId) REFERENCES authors(id) ON DELETE CASCADE,
                        FOREIGN KEY (articleId) REFERENCES articles(id) ON DELETE CASCADE
                    );
                """)

                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS journal_article (
                        JournalId INT,
                        ArticleId INT,
                        PRIMARY KEY (JournalId, ArticleId),
                        FOREIGN KEY (JournalId) REFERENCES journals(id) ON DELETE CASCADE,
                        FOREIGN KEY (ArticleId) REFERENCES articles(id) ON DELETE CASCADE
                    );
                """)

                print("Таблицы успешно созданы.")

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()

class EditAuthorWindow(QtWidgets.QMainWindow, Ui_EditAuthorWindow):
    def __init__(self, id_author):
        super(EditAuthorWindow, self).__init__()
        self.ui = Ui_EditAuthorWindow()
        self.ui.setupUi(self)
        self.author_id = id_author


        # Вызываем метод для запроса данный о пользоватеде по id из БД 
        info = self.get_info_about_author()

        # Устанавливаем стартовые значениея текущих данных из БД
        self.ui.edit_author_first_name.setText(str(info[1]))
        self.ui.edit_author_last_name.setText(str(info[2]))
        self.ui.edit_author_middle_name.setText(str(info[3]))
        self.ui.edit_author_info.setPlainText(str(info[4]))

        # Подключаем кнопку для сохранения(изменения) информации
        self.ui.putton_edit_author_save.clicked.connect(self.save_changes_info_author)

        # Установление фиксированного размеа окна
        self.setFixedSize(500, 380)
    
    # Метод для получения данных по текущем авторе
    def get_info_about_author(self):
        try:
            connection = mysql.connector.connect(
                host="sql12.freesqldatabase.com",
                user="sql12749774",
                password="kmYMIq9h6G",
                database="sql12749774"  # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                query = f"SELECT * FROM Authors WHERE ID_Author = %s;"
                cursor.execute(query, (self.author_id,))
    

                # Получаем все результаты при помощи fetchone
                authors = cursor.fetchone()
                print("Автор получен.")

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()

        return authors
        
    def save_changes_info_author(self):
        try:
            connection = mysql.connector.connect(
                host="sql12.freesqldatabase.com",
                user="sql12749774",
                password="kmYMIq9h6G",
                database="sql12749774"  # Имя базы данных
            )   

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()
                print(0)
                # Получаем данные из полей
                new_first_name = str(self.ui.edit_author_first_name.text())
                new_last_name = str(self.ui.edit_author_last_name.text())
                new_middle_name = str(self.ui.edit_author_middle_name.text())
                new_info = str(self.ui.edit_author_info.toPlainText())
                print(1)

                # Запрос на обновление данных
                query = """
                    UPDATE Authors 
                    SET first_name = %s, last_name = %s, middle_name = %s, info = %s 
                    WHERE ID_Author = %s
                """
                data = (new_first_name, new_last_name, new_middle_name, new_info, self.author_id)

                cursor.execute(query, data)
                connection.commit()  # Сохраняем изменения в базе данных

                print("Данные изменены.")

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()

class EditArticleWindow(QtWidgets.QMainWindow, Ui_EditArticleWindow):
    def __init__(self):
        super(EditArticleWindow, self).__init__()
        self.ui = Ui_EditArticleWindow()
        self.ui.setupUi(self)

class EditAuthorWidget(QWidget, Ui_Form):
    def __init__(self, main_window, id_author):
        super().__init__()
        self.ui = Ui_Form()  
        self.ui.setupUi(self)  

        self.main_window = main_window
        self.id_author = id_author
        
        # Подключаем нопку
        self.ui.putton_edit_author_save.clicked.connect(self.save_author_info)

        # self.ui.edit_author_first_name.setText(str(self.id_author))

        # Вызываем метод для запроса данный о пользоватеде по id из БД 
        info = self.get_info_about_author()

        # Устанавливаем стартовые значениея текущих данных из БД
        self.ui.edit_author_first_name.setText(str(info[1]))
        self.ui.edit_author_last_name.setText(str(info[2]))
        self.ui.edit_author_middle_name.setText(str(info[3]))
        self.ui.edit_author_info.setPlainText(str(info[4]))

        # Подключаем кнопку к методу
        # self.ui.someButton.clicked.connect(self.call_main_method)  # Замените 'someButton' на фактическое имя кнопки

    # Метод для сохранения изменений и обновления данных в списке авторов
    def save_author_info(self):

        # # Получаем список всех авторов, так как не передаем параметры
        # list_of_authors = self.get_list_of_authors()

        # # Передаем список авторов для добавления авторов на окно
        # self.loading_authors_from_the_database_to_page(list_of_authors)

        # Метод для изменени данных об авторе
        self.change_info_abot_author()

        # Получаем список всех авторов, так как не передаем параметры
        list_of_authors = self.main_window.get_list_of_authors()

        # Передаем список авторов для добавления авторов на окно
        self.main_window.loading_authors_from_the_database_to_page(list_of_authors)

        self.close()
    
    # Метод для получения данных о автора по его id
    def get_info_about_author(self):
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                query = f"SELECT * FROM Authors WHERE ID_Author = %s;"
                cursor.execute(query, (self.id_author,))
    

                # Получаем все результаты при помощи fetchone
                authors = cursor.fetchone()
                print("Автор получен.")

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()

        return authors

    # Метод для изменения данных автора в БД
    def change_info_abot_author(self):
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                # Получаем данные из полей
                new_first_name = str(self.ui.edit_author_first_name.text())
                new_last_name = str(self.ui.edit_author_last_name.text())
                new_middle_name = str(self.ui.edit_author_middle_name.text())
                new_info = str(self.ui.edit_author_info.toPlainText())

                # Запрос на обновление данных
                query = """
                    UPDATE Authors 
                    SET first_name = %s, last_name = %s, middle_name = %s, info = %s 
                    WHERE ID_Author = %s
                """
                data = (new_first_name, new_last_name, new_middle_name, new_info, self.id_author)

                cursor.execute(query, data)
                connection.commit()  # Сохраняем изменения в базе данных

                print("Данные изменены.")

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()

class EditArticleWidget(QWidget, Ui_Form_Article):
    def __init__(self, main_window, id_article):
        super().__init__()
        self.ui = Ui_Form_Article()  
        self.ui.setupUi(self)  

        # print(id_article)

        # Установление маски для ввода даты создания статьи
        self.ui.new_date_line_edit.setInputMask('0000-00-00;_')

        # Главное окно и id статьи
        self.main_window = main_window
        self.id_article = id_article
        
        # Подключаем нопку для сохранения изменений 
        self.ui.push_button_save_changes_article_info.clicked.connect(self.save_changes)

        # Получаем информацию и статье
        info_about_article = self.get_info_about_article_with_id()

        # Получаем id автора статьи
        self.id_author_article = info_about_article[0][0]
        # print(self.id_author_article)

        # Получаем всех авторов статей кроме текцщего
        all_authors = self.get_all_author()

        # Проверка получения данных
        # print(info_about_article)
        print("All authors = ")
        print(all_authors)

        current_article = info_about_article[0]

        # Данный из БД устанавливаем в окно
        self.ui.new_name_line_edit.setText(str(current_article[5]))
        self.ui.new_date_line_edit.setText(str(current_article[6]))
        self.ui.new_science_line_edit.setText(str(current_article[7]))
        self.ui.new_text_edit.setText(str(current_article[8]))

        # Установление авторов в combo box
        # Устанавливаем первым текущего автора статьи
        self.ui.new_author_combo_box.addItem(str(current_article[1]) + " " + str(current_article[2]) + " " + str(current_article[3]))
        # Добавляем остальных авторов
        for el in all_authors:
            self.ui.new_author_combo_box.addItem(str(el[1]) + " " + str(el[2]) + " " + str(el[3]))

    # Изменение информации о статье
    def change_info_about_article(self):
        new_name = self.ui.new_name_line_edit.text()
        new_data = self.ui.new_date_line_edit.text()
        new_science = self.ui.new_science_line_edit.text()
        new_text = self.ui.new_text_edit.toPlainText()
        new_name_author = self.ui.new_author_combo_box.currentText()
        # print("Автор = ",new_name_author)

        new_name_author = new_name_author.split()
        # print("Авто = ", new_name_author)

        new_name_author_first_name = new_name_author[0]
        new_name_author_last_name = new_name_author[1]
        new_name_author_middle_name = new_name_author[2]

        if not self.validate_date(new_data):
            QMessageBox.warning(self, "Ошибка", "Введите корректную дату в формате ГГГГ-ММ-ДД или РЕАЛЬНУЮ дату")
            self.ui.new_date_line_edit.clear()  # Очищаем поле ввода
            return
        elif not new_name:
            QMessageBox.warning(self, "Ошибка", "Введите название статьи")
            return
        elif not new_science:
            QMessageBox.warning(self, "Ошибка", "Введите название науки")
            return
        elif not new_text:
            QMessageBox.warning(self, "Ошибка", "Введите название текст")
            return
        elif len(new_name) > 255:
            QMessageBox.warning(self, "Ошибка", "Введите название статьи меньше")
            return
        elif len(new_science) > 255:
            QMessageBox.warning(self, "Ошибка", "Введите название статьи меньше")
            return

        # Изменяем информацию о статье
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                # Формируем SQL-запрос по параметрам
                query = f"UPDATE Articles SET name = %s, date_create = %s, science = %s, text_article = %s WHERE ID_Article = %s;"
                cursor.execute(query, (new_name, new_data, new_science, new_text, self.id_article,))
                print("Данные обнавлены")
                connection.commit()

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()
        

        # Меняем связь между автором и статьёй
        # Получаем нового автора по его имени и фамилии
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                query = f"SELECT ID_Author FROM Authors WHERE first_name = %s AND last_name = %s AND middle_name = %s;"
                cursor.execute(query, (new_name_author_first_name, new_name_author_last_name, new_name_author_middle_name,))
                id_new_author = cursor.fetchall()
                # print("NEW ID AUTHRO = ")
                # print(id_new_author)
                print("Данные обнавлены")
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()

        print(f"ID article = {self.id_article}")
        print(f"ID old author = {self.id_author_article}")
        print(f"ID new author = {id_new_author[0][0]}")

        self.change_connect_author_with_article(self.id_article, self.id_author_article, id_new_author[0][0])

    # Обновление соедниения автора с статьей
    def change_connect_author_with_article(self, id_article, old_id_author, new_id_author):

        print("Param = ")
        print(id_article)
        print(old_id_author)
        print(new_id_author)

        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                query = f"DELETE FROM Authors_Articles WHERE ID_Article = %s;"
                cursor.execute(query, (id_article,))

                connection.commit()

                query = f"INSERT INTO Authors_Articles (ID_Author, ID_Article) VALUES (%s, %s);"
                cursor.execute(query, (new_id_author, id_article,))

                connection.commit()

                print("Данные обнавлены")
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()     

    def save_changes(self):

        self.change_info_about_article()

        # Получение всех статей
        list_articles = self.main_window.get_all_articles_from_db()

        # Загрузка статей в widget с статьями
        self.main_window.load_articles_from_db_to_page_articles(list_articles)

        self.close()
        
        # print(self.id_article)
    
    # Получение информации о статье
    def get_info_about_article_with_id(self):
        
        # print("ID")
        # print(self.id_article)

        onnection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                cursor.execute("""
                SELECT
                    Authors.ID_Author,
                    Authors.first_name,
                    Authors.last_name,
                    Authors.middle_name,
                    Articles.ID_Article,
                    Articles.name,
                    Articles.date_create,
                    Articles.science,
                    Articles.text_article
                FROM
                    Articles
                INNER JOIN
                    Authors_Articles ON Articles.ID_Article = Authors_Articles.ID_Article
                INNER JOIN
                    Authors ON Authors_Articles.ID_Author = Authors.ID_Author
                WHERE
                    Articles.ID_Article = %s;
                """, (self.id_article,))
                
                items = cursor.fetchall()
                # print(items)
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        
        return items

    # Получение всех авторов
    def get_all_author(self):
        onnection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7751998",
                password="7kPDaYU2TX",
                database="sql7751998" # Имя базы данных
            )

            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()

                cursor.execute("""
                SELECT ID_Author, first_name, last_name, middle_name FROM Authors WHERE ID_Author <> %s;
                """, (self.id_author_article,))
                
                items = cursor.fetchall()
                # print(items)
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        
        return items

        # Проверка корректоности даты на странице статей 
    
    # Проверка корректности даты
    def validate_date(self, date_string):
        # Проверяем, что строка имеет длину 10 символов (формат ГГГГ-ММ-ДД)
        if len(date_string) != 10:
            return False

        year = date_string[:4]
        month = date_string[5:7]
        day = date_string[8:]

        # Проверка даты
        # print(year)
        # print(month)
        # print(day)

        # Проверка года
        if int(year) >= 1000:
            # print(1)
            # Проверка месяца
            if int(month) >= 1 and int(month) <= 12:
                # Проверка дня
                if int(day) >= 1 and int(day) <= 31:
                    # Дополнительная проверка на количество дней в месяце
                    if month in ['04', '06', '09', '11'] and int(day) == 31:
                        return False  # Апрель, Июнь, Сентябрь, Ноябрь имеют максимум 30 дней
                    if month == '02':
                        # Проверка на високосный год
                        if (int(year) % 4 == 0 and int(year) % 100 != 0) or (int(year) % 400 == 0):
                            if int(day) > 29:
                                return False
                        else:
                            if int(day) > 28:
                                return False
                    return True
        return False

if __name__ == "__main__":

    app = QtWidgets.QApplication(sys.argv)
    window = MainWindow()
    window.showFullScreen()
    sys.exit(app.exec_())

    # app = QtWidgets.QApplication(sys.argv)
    # window = EditAuthorWindow()
    # window.show()
    # sys.exit(app.exec_())

    # app = QtWidgets.QApplication(sys.argv)
    # window = EditArticleWindow()
    # window.show()
    # sys.exit(app.exec_())

