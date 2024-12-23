import sys
import os
from PyQt5 import QtWidgets
from PyQt5.QtWidgets import QMessageBox, QListWidgetItem, QVBoxLayout, QWidget, QLabel, QPushButton, QHBoxLayout
from PyQt5 import QtGui
from datetime import datetime
from docx import Document
import aspose.pdf as ap
from MainWindow import Ui_MainWindow
from EditAuthorWidget import Ui_Form  
from EditArticleWidget import Ui_Form_Article
import mysql.connector
from mysql.connector import connect, Error

class MainWindow(QtWidgets.QMainWindow, Ui_MainWindow):
    def __init__(self):
        super(MainWindow, self).__init__()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.create_db()
        self.open_page_of_authors()
        self.ui.button_exit.clicked.connect(self.close_app)
        self.ui.button_open_page_author.clicked.connect(self.open_page_of_authors)
        self.ui.button_open_page_articles.clicked.connect(self.open_page_of_articles)
        self.ui.button_open_page_journals.clicked.connect(self.open_page_of_journals)

        self.ui.button_open_page_author.setStyleSheet("""
            QPushButton {
                background-color: #161a1d;
                color: white;
                border-radius: 5px;
                padding: 7px 10px 7px 10px;
                font-size: 15px;
            }
            QPushButton:pressed {
                background-color: #343a40;  /* Цвет при нажатии */
            }
        """)

        self.ui.button_open_page_articles.setStyleSheet("""
            QPushButton {
                background-color: #161a1d;
                color: white;
                border-radius: 5px;
                padding: 7px 10px 7px 10px;
                font-size: 15px;
            }
            QPushButton:pressed {
                background-color: #343a40;  /* Цвет при нажатии */
            }
        """)

        self.ui.button_open_page_journals.setStyleSheet("""
            QPushButton {
                background-color: #161a1d;
                color: white;
                border-radius: 5px;
                padding: 7px 10px 7px 10px;
                font-size: 15px;
            }
            QPushButton:pressed {
                background-color: #343a40;  /* Цвет при нажатии */
            }
        """)

        self.add_to_combo_box_authors()    
        self.add_to_combo_box_articles()
        self.add_to_combo_box_journals()
        self.ui.button_add_author.clicked.connect(self.add_author_to_db)
        self.ui.button_search_authors_with_param.clicked.connect(self.search_authors_with_param)
        self.ui.button_add_articles.clicked.connect(self.add_article_to_db)
        self.ui.button_search_articles_with_param.clicked.connect(self.search_articles_with_param)
        self.ui.push_button_add_journals.clicked.connect(self.add_journal_to_db)
        self.ui.button_search_journals_with_param.clicked.connect(self.search_journals_with_param)
        self.ui.combo_box_journals.currentIndexChanged.connect(self.set_mask_to_line_edit_value_param_journals)

    def read_db_config(self, filename):
        config = {}
        with open(filename, 'r') as file:
            for line in file:
                key, value = line.strip().split('=')
                config[key] = value
        return config
    def set_mask_to_line_edit_value_param_journals(self):
        if self.ui.combo_box_journals.currentText() == "Дата":
            self.ui.line_edit_value_param_journals.setInputMask('0000-00-00;_')
        else:
            self.ui.line_edit_value_param_journals.setInputMask('')

    def close_app(self):
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                cursor = connection.cursor()
                query = """SELECT * FROM Journals;"""
                self.write_sql(query)
                cursor.execute(query)
                items = cursor.fetchall()
                query = """
                        SELECT 
                            Journals.ID_Journal,
                            Articles.ID_Article,
	                        Articles.name,
                            Articles.science,
                            Articles.date_create,
                            Articles.text_article 
                        FROM 
                            Journals 
                        INNER JOIN 
                            Articles_Journals ON Articles_Journals.ID_Journal = Journals.ID_Journal 
                        INNER JOIN 
                            Articles ON Articles.ID_Article = Articles_Journals.ID_Article;
                    """
                cursor.execute(query)
                info_journal = cursor.fetchall()

        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()
    
        count_journals = len(items)
        doc = Document()
        doc.add_heading(f"Отчёт о системе")
        doc.add_paragraph("   ")
        doc.add_paragraph(f"Количество журналов: {count_journals}")
        doc.add_paragraph("   ")
        for item in items:
            id_article = item[0]
            doc.add_paragraph(f"Журнал : {item[1]} {item[2]}")
            doc.add_paragraph("   ")
            for article in info_journal:
                if article[0] == id_article:
                    doc.add_paragraph(f"Название статьи : {str(article[2])}")
                    doc.add_paragraph(f"Наука : {str(article[3])}")
                    doc.add_paragraph(f"Дата : {str(article[4])}")
                    doc.add_paragraph(f"Текст : {str(article[5])}")
                    doc.add_paragraph("   ")
            doc.add_paragraph("   ")
            doc.add_paragraph("   ")

        project_path = os.path.dirname(os.path.abspath(sys.modules['__main__'].__file__))
        file_name = f'report.docx'
        file_path = os.path.join(project_path, file_name)
        doc.save(file_path)
        self.close()

    def open_page_of_authors(self):
        self.ui.Main_widgets_pages.setCurrentIndex(0)
        list_of_authors = self.get_list_of_authors()
        self.loading_authors_from_the_database_to_page(list_of_authors)

    def open_page_of_articles(self):
        self.ui.Main_widgets_pages.setCurrentIndex(2)
        self.load_authors_to_combo_box_article()
        self.ui.line_edit_data_article.setInputMask('0000-00-00;_')
        list_articles = self.get_all_articles_from_db()
        self.load_articles_from_db_to_page_articles(list_articles)

    def open_page_of_journals(self):
        self.ui.Main_widgets_pages.setCurrentIndex(1)
        self.ui.line_edit_journals_data.setInputMask('0000-00-00;_')
        list_of_articles_no_use = self.loading_all_articles_with_not_use()
        self.insert_articles_to_widget_no_use(list_of_articles_no_use)
        all_journals = self.get_all_journals_from_db()
        self.loading_journasl_to_list_widget_page_journals(all_journals)

    def add_to_combo_box_authors(self):
        self.ui.combo_box_authors.addItem("Имя")
        self.ui.combo_box_authors.addItem("Фамилия")
        self.ui.combo_box_authors.addItem("Отчество")
        self.ui.combo_box_authors.addItem("Дополнительная информация")
        self.ui.combo_box_authors.addItem("Сброс параметров")
    def add_to_combo_box_articles(self):

        self.ui.combo_box_articles.addItem("Название статьи")
        self.ui.combo_box_articles.addItem("Наука")
        self.ui.combo_box_articles.addItem("Текст статьи")
        self.ui.combo_box_articles.addItem("Сброс параметров")

    def add_to_combo_box_journals(self):
        self.ui.combo_box_journals.addItem("Название")
        self.ui.combo_box_journals.addItem("Дата")
        self.ui.combo_box_journals.addItem("Номер")
        self.ui.combo_box_journals.addItem("Сброс параметров")

    def add_author_to_db(self):
        first_name = self.ui.author_first_name.text()
        last_name = self.ui.author_last_name.text()
        middle_name = self.ui.author_middle_name.text()
        info_about_author = self.ui.author_additional_info.toPlainText()
        errors = []       
        if not first_name:
            errors.append('Необходимо указать имя автора.')
        if not last_name:
            errors.append('Необходимо указать фамилию автора.')
        if not middle_name:
            errors.append('Необходимо указать отчество автора.')
        if not info_about_author:
            errors.append('Необходимо ввести дополнительные данные об авторе.')
        if len(first_name) < 2:
            errors.append(f'Имя не может состоять из {len(first_name)} символа (символов).')
        if len(last_name) < 2:
            errors.append(f'Фамилия не может состоять из {len(last_name)} символа (символов).')
        if len(middle_name) < 5:
            errors.append(f'Отчество не может состоять из {len(middle_name)} символа (символов).')
        if errors:
            message_box = QMessageBox()
            message_box.setIcon(QMessageBox.Warning)
            message_box.setWindowTitle("Ошибка")
            message_box.setText("\n".join(errors))
            message_box.exec_()
            return
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()
                cursor.execute("""INSERT INTO Authors (first_name, last_name, middle_name, info)
                VALUES (%s, %s, %s, %s); """, (first_name, last_name, middle_name, info_about_author))
                query = f"""INSERT INTO Authors (first_name, last_name, middle_name, info) VALUES ({first_name}, {last_name}, {middle_name}, {info_about_author}); """
                self.write_sql(query)
                print("Пользователь добавлен")
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        self.ui.author_first_name.clear()
        self.ui.author_last_name.clear()
        self.ui.author_middle_name.clear()
        self.ui.author_additional_info.clear()
        list_of_authors = self.get_list_of_authors()
        self.loading_authors_from_the_database_to_page(list_of_authors)

    def loading_authors_from_the_database_to_page(self, list_of_authors):
        self.ui.Widget_authors.clear()
        for author in list_of_authors:
            item = QListWidgetItem()
            item_widget = QWidget()
            id_author = QLabel(str(author[0]))
            separator_id = QLabel()
            name_author = QLabel(str(author[1]) + " " + str(author[2]) + " " + str(author[3]))
            separator_name = QLabel()
            if len(str(author[4])) > 95:
                text_info = QLabel(str(author[4][:95]) + "... ")
            else:
                text_info = QLabel(str(author[4]))
            separator_info = QLabel()
            edit_push_button = QPushButton("Редактировать")
            delete_push_button = QPushButton("Удалить")
            edit_push_button.setObjectName(str(author[0]))
            delete_push_button.setObjectName(str(author[0])) 
            edit_push_button.clicked.connect(self.edit_author_push_button)
            delete_push_button.clicked.connect(self.delete_author_push_button)
            edit_push_button.setStyleSheet("""
                QPushButton {
                    background-color: #d90429; 
                    color: white; 
                    border-radius: 5px; 
                    font-size: 17px; 
                    padding: 4px 7px;
                }
                QPushButton:pressed {
                    background-color: #ff5a5f;  /* Цвет при нажатии */
                }
            """)
            delete_push_button.setStyleSheet("""
                QPushButton {
                    background-color: #d90429; 
                    color: white; 
                    border-radius: 5px; 
                    font-size: 17px; 
                    padding: 4px 7px;
                }
                QPushButton:pressed {
                    background-color: #ff5a5f;  /* Цвет при нажатии */
                }
            """)
            item_layout = QHBoxLayout()
            item_layout.addWidget(id_author)
            item_layout.addWidget(separator_id)
            item_layout.addWidget(name_author)
            item_layout.addWidget(separator_name)
            item_layout.addWidget(text_info)
            item_layout.addWidget(separator_info)
            item_layout.addWidget(edit_push_button)
            item_layout.addWidget(delete_push_button)
            item_widget.setLayout(item_layout)
            item_widget.setStyleSheet("background-color: #81c3d7; font-size: 17px; border-radius: 10px; margin-bottom: 9px; margin-top: 9px; margin-right: 5px; margin-left: 5px; ")  
            item.setSizeHint(item_widget.sizeHint())
            self.ui.Widget_authors.addItem(item)
            self.ui.Widget_authors.setItemWidget(item, item_widget)

    def get_list_of_authors(self, param=None, value=None):
        connection = None
        cursor = None
        authors = []
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                cursor = connection.cursor()
                if param == "info":
                    query = f"SELECT * FROM Authors WHERE {param} LIKE %s;"
                    cursor.execute(query, (f'%{value}%',))
                elif param and value:
                    query = f"SELECT * FROM Authors WHERE {param} = %s;"
                    cursor.execute(query, (value,))
                else:
                    cursor.execute("SELECT * FROM Authors;")
                authors = cursor.fetchall()
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()
        return authors

    def edit_author_push_button(self):
        sender = self.sender()
        push_button = self.findChild(QPushButton, sender.objectName())
        self.edit_author_widget = EditAuthorWidget(self, str(push_button.objectName()))
        self.edit_author_widget.show()

    def delete_author_push_button(self):
        sender = self.sender()
        push_button = self.findChild(QPushButton, sender.objectName())
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                cursor = connection.cursor()
                query = f"SELECT ID_Article FROM Authors_Articles WHERE ID_Author = %s;"
                cursor.execute(query, (push_button.objectName(),))
                query = f"""SELECT ID_Article FROM Authors_Articles WHERE ID_Author = {push_button.objectName()};"""
                self.write_sql(query)
                id_articles = cursor.fetchall()
                print(id_articles)
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        id_for_delete = []
        for el in id_articles:
            id_for_delete.append(el[0])
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()
                for el in id_for_delete:
                    id_article = el 
                    query = f"DELETE FROM Articles WHERE ID_Article = %s;"
                    cursor.execute(query, (id_article,))
                    query = f"""DELETE FROM Articles WHERE ID_Article = {id_article};"""
                    self.write_sql(query)
                print("Статьи удалены.")
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()
                query = f"DELETE FROM Authors WHERE ID_Author = %s;"
                cursor.execute(query, (push_button.objectName(),))
                query = f"""DELETE FROM Authors WHERE ID_Author = {push_button.objectName()};"""
                self.write_sql(query)
                print(f"Пользователь удалён. {push_button.objectName()}")
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        list_of_authors = self.get_list_of_authors()
        self.loading_authors_from_the_database_to_page(list_of_authors)

    def search_authors_with_param(self):
        value_text = self.ui.line_edit_author_param.text()
        value_param = self.ui.combo_box_authors.currentText()
        if value_param == "Имя":
            current_param = "first_name"
        elif value_param == "Фамилия":
            current_param = "last_name"
        elif value_param == "Отчество":
            current_param = "middle_name"
        elif value_param == "Дополнительная информация":
            current_param = "info"
        elif value_param == "Сброс параметров":
            value_param = None
            current_param = None
            self.ui.line_edit_author_param.clear()
        info = self.get_list_of_authors(current_param, value_text)
        self.loading_authors_from_the_database_to_page(info)
    
    def search_articles_with_param(self):
        param = self.ui.combo_box_articles.currentText()
        value = self.ui.line_edit_articles_param.text()
        if param == "Название статьи":
            current_param = "name"
        elif param == "Наука":
            current_param = "science"
        elif param == "Текст статьи":
            current_param = "text_article"
        elif param == "Сброс параметров": 
            param = None
            current_param = None
            self.ui.line_edit_author_param.clear()
        list_of_articles_with_param = self.loading_articles_from_db_with_param(current_param, value)
        self.load_articles_from_db_to_page_articles(list_of_articles_with_param)

    def loading_articles_from_db_with_param(self, param=None, value=None):
        connection = None
        cursor = None
        authors = []
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                cursor = connection.cursor()
                if param == "text_article":
                    query = f"""
                        SELECT
                            Authors.ID_Author,
                            Authors.first_name,
                            Authors.last_name,
                            Articles.ID_Article,
                            Articles.name,
                            Articles.date_create,
                            Articles.science,
                            Articles.text_article
                        FROM
                            Articles
                        INNER JOIN
                            Authors_Articles ON Articles.ID_Article = Authors_Articles.ID_Article
                        INNER JOIN
                            Authors ON Authors_Articles.ID_Author = Authors.ID_Author
                        WHERE 
                            Articles.{param} LIKE %s;
                        """
                    cursor.execute(query, (f'%{value}%',))
                    query = f""" SELECT
                            Authors.ID_Author,
                            Authors.first_name,
                            Authors.last_name,
                            Articles.ID_Article,
                            Articles.name,
                            Articles.date_create,
                            Articles.science,
                            Articles.text_article
                        FROM
                            Articles
                        INNER JOIN
                            Authors_Articles ON Articles.ID_Article = Authors_Articles.ID_Article
                        INNER JOIN
                            Authors ON Authors_Articles.ID_Author = Authors.ID_Author
                        WHERE 
                            Articles.{param} = %{value}%;"""
                    self.write_sql(query)  
                elif param and value:
                    query = f"""
                        SELECT
                            Authors.ID_Author,
                            Authors.first_name,
                            Authors.last_name,
                            Articles.ID_Article,
                            Articles.name,
                            Articles.date_create,
                            Articles.science,
                            Articles.text_article
                        FROM
                            Articles
                        INNER JOIN
                            Authors_Articles ON Articles.ID_Article = Authors_Articles.ID_Article
                        INNER JOIN
                            Authors ON Authors_Articles.ID_Author = Authors.ID_Author
                        WHERE 
                            Articles.{param} = %s;
                        """
                    cursor.execute(query, (value,))
                    query = f""" SELECT
                            Authors.ID_Author,
                            Authors.first_name,
                            Authors.last_name,
                            Articles.ID_Article,
                            Articles.name,
                            Articles.date_create,
                            Articles.science,
                            Articles.text_article
                        FROM
                            Articles
                        INNER JOIN
                            Authors_Articles ON Articles.ID_Article = Authors_Articles.ID_Article
                        INNER JOIN
                            Authors ON Authors_Articles.ID_Author = Authors.ID_Author
                        WHERE 
                            Articles.{param} = {value};"""
                    self.write_sql(query)
                else:
                    query = f"""
                        SELECT
                            Authors.ID_Author,
                            Authors.first_name,
                            Authors.last_name,
                            Articles.ID_Article,
                            Articles.name,
                            Articles.date_create,
                            Articles.science,
                            Articles.text_article
                        FROM
                            Articles
                        INNER JOIN
                            Authors_Articles ON Articles.ID_Article = Authors_Articles.ID_Article
                        INNER JOIN
                            Authors ON Authors_Articles.ID_Author = Authors.ID_Author;
                        """
                    cursor.execute(query)
                    self.write_sql(query)
                authors = cursor.fetchall()
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()
        return authors


    def add_article_to_db(self):
        if self.ui.combo_box_authors_article.count() == 0:
            QMessageBox.warning(self, "Ошибка", "Выберите автора")
            return 
        text = self.ui.combo_box_authors_article.currentText()
        arr_info_about_current_author = text.split()
        id_author = arr_info_about_current_author[0]
        first_name = arr_info_about_current_author[1]
        last_name = arr_info_about_current_author[2]
        middle_name = arr_info_about_current_author[3]
        
        id_author = str(id_author)
        first_name = str(first_name)
        last_name = str(last_name)
        middle_name = str(middle_name)
        name_articel = self.ui.line_edit_name_article.text()
        data_articel = self.ui.line_edit_data_article.text() 
        science_articel = self.ui.line_edit_scince_article.text()
        text_article = self.ui.text_edit_text_article.toPlainText()
        id_author_for_articel = self.get_author_id_with_first_last_middle_name(id_author, first_name, last_name, middle_name)
        if not self.validate_date(data_articel):
            QMessageBox.warning(self, "Ошибка", "Введите корректную дату в формате ГГГГ-ММ-ДД или РЕАЛЬНУЮ дату")
            self.ui.line_edit_data_article.clear()
            return
        elif not name_articel:
            QMessageBox.warning(self, "Ошибка", "Введите навзание статьи")
            return
        elif not science_articel:
            QMessageBox.warning(self, "Ошибка", "Введите навзание науки")
            return
        elif not text_article:
            QMessageBox.warning(self, "Ошибка", "Введите текст статьи")
            return
        elif len(name_articel) >= 255:
            QMessageBox.warning(self, "Ошибка", "Введите навзание стать меньеше 255 символов")
            return
        elif len(science_articel) >= 255:
            QMessageBox.warning(self, "Ошибка", "Введите навзание науки меньше 255 символов")
            return
        connection = None
        cursor = None

        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735"
            )
            if connection.is_connected():
                cursor = connection.cursor()
                print(1)
                cursor.execute("""INSERT INTO Articles (name, date_create, science, text_article)
                VALUES (%s, %s, %s, %s); """, (name_articel, data_articel, science_articel, text_article))
                query = f"""INSERT INTO Articles (name, date_create, science, text_article) VALUES ({name_articel}, {data_articel}, {science_articel}, {text_article});"""
                self.write_sql(query)
        except Error as e:
            print(f"Ошибка подключения: {e}")
            QMessageBox.warning(self, "Ошибка", "Название статьи должно быть уникальным")
            return
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        self.ui.line_edit_name_article.clear()
        self.ui.line_edit_data_article.clear()
        self.ui.line_edit_scince_article.clear()
        self.ui.text_edit_text_article.clear()

        id_article = self.get_id_article_with_name_date_science_text(name_articel, data_articel, science_articel, text_article)
        self.connect_author_and_article_in_db(id_author_for_articel, id_article)
        list_articles = self.get_all_articles_from_db()

        self.load_articles_from_db_to_page_articles(list_articles)

    def load_authors_to_combo_box_article(self):
        self.ui.combo_box_authors_article.clear()
        all_authors = self.get_list_of_authors()
        arr_authors = []
        for author in all_authors:
            el_author = str(author[0]) + " " + author[1] + " " + author[2] + " " + author[3]
            arr_authors.append(el_author)
        self.ui.combo_box_authors_article.addItems(arr_authors)        

    def get_author_id_with_first_last_middle_name(self, id_a, first, last, middle):
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735"
            )
            if connection.is_connected():
                cursor = connection.cursor()
                query = f"SELECT * FROM Authors WHERE ID_Author = %s AND first_name = %s AND last_name = %s AND middle_name = %s"
                cursor.execute(query, (id_a, first, last, middle),)
                query = f"""SELECT * FROM Authors WHERE ID_Author = {id_a} AND first_name = {first} AND last_name = {last} AND middle_name = {middle}"""
                self.write_sql(query)
                authors = cursor.fetchone()
                author_id = authors[0]
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()

        return author_id

    def get_id_article_with_name_date_science_text(self, name, date, science, text):
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                cursor = connection.cursor()
                cursor.execute("""SELECT * FROM Articles WHERE name = %s AND date_create = %s AND science = %s AND text_article = %s""", (name, date, science, text))
                articles = cursor.fetchone()
                articles_id = articles[0]
                query = f"""SELECT * FROM Articles WHERE name = {name} AND date_create = {date} AND science = {science} AND text_article = {text};"""
                self.write_sql(query)
                return articles_id
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
    def connect_author_and_article_in_db(self, id_author, id_article):
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()
                cursor.execute("""INSERT INTO Authors_Articles (ID_Author,  ID_Article) 
                VALUES (%s, %s); """, (id_author, id_article))
                query = f"""INSERT INTO Authors_Articles (ID_Author,  ID_Article) VALUES ({id_author}, {id_article}); """
                self.write_sql(query)
                print("Автор и Статья подключены")
                cursor.execute("""UPDATE Authors SET count_articles = count_articles + 1 WHERE ID_Author = %s; """, (id_author, ))
                query = f"""UPDATE Authors SET count_articles = count_articles + 1 WHERE ID_Author = {id_author};"""
                self.write_sql(query)
                print("Количество статей увеличено")
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()

    def load_articles_from_db_to_page_articles(self, list_of_articles):
        self.ui.list_widget_articles.clear()
        for article in list_of_articles:
            item = QListWidgetItem()
            item_widget = QWidget()
            name_author = QLabel(str(article[1]) + " " + str(article[2]))
            name_article = QLabel(str(article[4]))
            date_create_article = QLabel(str(article[5]))
            science_of_article = QLabel(str(article[6]))
            text_of_article = str(article[7])
            if len(text_of_article) > 35:
                text_of_article = QLabel(str(article[7][:35]) + "...")
            else:
                text_of_article = QLabel(str(article[7]))
            edit_push_button = QPushButton("Редактировать")
            delete_push_button = QPushButton("Удалить")
            view_push_button = QPushButton("Смотреть")
            edit_push_button.setObjectName(str(article[3])) #
            delete_push_button.setObjectName(str(article[3]))  
            view_push_button.setObjectName(str(article[3]))
            edit_push_button.clicked.connect(self.edit_article_push_button)
            delete_push_button.clicked.connect(self.delete_article_push_button)
            view_push_button.clicked.connect(self.view_article_push_button)
            edit_push_button.setStyleSheet("""
                QPushButton {
                    background-color: #d90429; 
                    color: white; 
                    border-radius: 5px; 
                    font-size: 17px; 
                    padding: 4px 7px;
                }
                QPushButton:pressed {
                    background-color: #ff5a5f;  /* Цвет при нажатии */
                }
            """)
            view_push_button.setStyleSheet("""
                QPushButton {
                    background-color: #d90429; 
                    color: white; 
                    border-radius: 5px; 
                    font-size: 17px; 
                    padding: 4px 7px;
                }
                QPushButton:pressed {
                    background-color: #ff5a5f;  /* Цвет при нажатии */
                }
            """)
            delete_push_button.setStyleSheet("""
                QPushButton {
                    background-color: #d90429; 
                    color: white; 
                    border-radius: 5px; 
                    font-size: 17px; 
                    padding: 4px 7px;
                }
                QPushButton:pressed {
                    background-color: #ff5a5f;  /* Цвет при нажатии */
                }
            """)
            item_layout = QHBoxLayout()
            item_layout.addWidget(name_author)
            item_layout.addWidget(name_article)
            item_layout.addWidget(date_create_article)
            item_layout.addWidget(science_of_article)
            item_layout.addWidget(text_of_article)
            item_layout.addWidget(edit_push_button)
            item_layout.addWidget(delete_push_button)
            item_layout.addWidget(view_push_button)
            item_widget.setLayout(item_layout)
            item_widget.setStyleSheet("background-color: #81c3d7; font-size: 17px; border-radius: 10px; margin-bottom: 9px; margin-top: 9px; margin-right: 5px; margin-left: 5px; ")  
            item.setSizeHint(item_widget.sizeHint())
            self.ui.list_widget_articles.addItem(item)
            self.ui.list_widget_articles.setItemWidget(item, item_widget)

    def edit_article_push_button(self):
        sender = self.sender()
        push_button = self.findChild(QPushButton, sender.objectName())
        self.edit_article_widget = EditArticleWidget(self, str(push_button.objectName()))
        self.edit_article_widget.show()

    def delete_article_push_button(self):
        sender = self.sender()
        push_button = self.findChild(QPushButton, sender.objectName())
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                cursor = connection.cursor()
                cursor.execute("""DELETE FROM Articles WHERE ID_Article = %s """, (push_button.objectName(),))
                query = f"""DELETE FROM Articles WHERE ID_Article = {push_button.objectName()};"""
                self.write_sql(query)
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        list_articles = self.get_all_articles_from_db()
        self.load_articles_from_db_to_page_articles(list_articles)

    def view_article_push_button(self):
        sender = self.sender()
        push_button = self.findChild(QPushButton, sender.objectName())
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()
                query = """
                SELECT
                    Authors.ID_Author,
                    Authors.first_name,
                    Authors.last_name,
                    Articles.ID_Article,
                    Articles.name,
                    Articles.date_create,
                    Articles.science,
                    Articles.text_article
                FROM
                    Articles
                INNER JOIN
                    Authors_Articles ON Articles.ID_Article = Authors_Articles.ID_Article
                INNER JOIN
                    Authors ON Authors_Articles.ID_Author = Authors.ID_Author
                WHERE
                    Articles.ID_Article = %s;
                """
                cursor.execute(query, (push_button.objectName(),))
                query = f"""
                SELECT
                    Authors.ID_Author,
                    Authors.first_name,
                    Authors.last_name,
                    Articles.ID_Article,
                    Articles.name,
                    Articles.date_create,
                    Articles.science,
                    Articles.text_article
                FROM
                    Articles
                INNER JOIN
                    Authors_Articles ON Articles.ID_Article = Authors_Articles.ID_Article
                INNER JOIN
                    Authors ON Authors_Articles.ID_Author = Authors.ID_Author
                WHERE
                    Articles.ID_Article = {push_button.objectName()};"""
                self.write_sql(query)
                items = cursor.fetchone()
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        doc = Document()
        doc.add_heading(f"{items[4]}")
        doc.add_paragraph(f"Автор: {items[1]} {items[2]}")
        doc.add_paragraph(f"Дата: {items[5]}")
        doc.add_paragraph(f"Наука: {items[6]}")
        doc.add_paragraph(f"{items[7]}")
        project_path = os.path.dirname(os.path.abspath(sys.modules['__main__'].__file__))
        file_name = f'article_{push_button.objectName()}.docx'
        file_path = os.path.join(project_path, file_name)
        doc.save(file_path)
        os.startfile(file_path)

    def get_all_articles_from_db(self):
        onnection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                cursor = connection.cursor()
                cursor.execute("""
                SELECT
                    Authors.ID_Author,
                    Authors.first_name,
                    Authors.last_name,
                    Articles.ID_Article,
                    Articles.name,
                    Articles.date_create,
                    Articles.science,
                    Articles.text_article
                FROM
                    Articles
                INNER JOIN
                    Authors_Articles ON Articles.ID_Article = Authors_Articles.ID_Article
                INNER JOIN
                    Authors ON Authors_Articles.ID_Author = Authors.ID_Author;
                """)
                query = f"""
                                SELECT
                    Authors.ID_Author,
                    Authors.first_name,
                    Authors.last_name,
                    Articles.ID_Article,
                    Articles.name,
                    Articles.date_create,
                    Articles.science,
                    Articles.text_article
                FROM
                    Articles
                INNER JOIN
                    Authors_Articles ON Articles.ID_Article = Authors_Articles.ID_Article
                INNER JOIN
                    Authors ON Authors_Articles.ID_Author = Authors.ID_Author;"""
                self.write_sql(query)
                items = cursor.fetchall()
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        return items

    def validate_date(self, date_string):
        if len(date_string) != 10:
            return False
        year = date_string[:4]
        month = date_string[5:7]
        day = date_string[8:]
        if int(year) >= 1000:
            if int(month) >= 1 and int(month) <= 12:
                if int(day) >= 1 and int(day) <= 31:
                    if month in ['04', '06', '09', '11'] and int(day) == 31:
                        return False  
                    if month == '02':
                        if (int(year) % 4 == 0 and int(year) % 100 != 0) or (int(year) % 400 == 0):
                            if int(day) > 29:
                                return False
                        else:
                            if int(day) > 28:
                                return False
                    return True
        return False

    def add_journal_to_db(self):
        name_journal = self.ui.line_edit_journals_name.text()
        date_journal = self.ui.line_edit_journals_data.text()
        number_journal = self.ui.line_edit_journals_number.text()
        if not self.validate_date(date_journal):
            QMessageBox.warning(self, "Ошибка", "Введите корректную дату в формате ГГГГ-ММ-ДД или РЕАЛЬНУЮ дату")
            self.ui.line_edit_journals_data.clear()
            return
        elif len(name_journal) >= 255:
            QMessageBox.warning(self, "Ошибка", "Введите название журнала поменьше")
            return
        elif not name_journal:
            QMessageBox.warning(self, "Ошибка", "Введите название журнала")
            return
        elif not number_journal.isdigit():
            QMessageBox.warning(self, "Ошибка", "Номер журнала должен быть числом")
            return
        elif int(number_journal) >= 2147483647:
            QMessageBox.warning(self, "Ошибка", "Введите число поменьше")
            return

        if self.add_journals_to_db(name_journal, date_journal, number_journal):
            return
        self.ui.line_edit_journals_name.clear()
        self.ui.line_edit_journals_data.clear()
        self.ui.line_edit_journals_number.clear()
        all_journals = self.get_all_journals_from_db()
        self.loading_journasl_to_list_widget_page_journals(all_journals)
        id_journal = self.get_id_journal(name_journal, date_journal, number_journal)
        id_journal = id_journal[0]
        list_name = []
        item_count = self.ui.list_wodget_use_articles.count()
        for index in range(item_count):
            item = self.ui.list_wodget_use_articles.item(index)  
            if item is not None:
                article_name = item.text() 
                list_name.append(article_name)
        list_id_articles = self.get_id_articles(list_name)
        self.connect_article_with_journal(id_journal, list_id_articles)
        self.change_status_articles(list_id_articles)
        all_journals = self.get_all_journals_from_db()
        self.loading_journasl_to_list_widget_page_journals(all_journals)
        list_of_articles_no_use = self.loading_all_articles_with_not_use()
        self.insert_articles_to_widget_no_use(list_of_articles_no_use)
        self.ui.list_wodget_use_articles.clear()

    def loading_all_articles_with_not_use(self):
        onnection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                cursor = connection.cursor()
                cursor.execute("""
                SELECT ID_Article, name FROM Articles WHERE isUse = 0;
                """)
                query = f"""SELECT ID_Article, name FROM Articles WHERE isUse = 0;"""
                self.write_sql(query)
                list_of_articles_no_use = cursor.fetchall()
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        
        return list_of_articles_no_use

    def insert_articles_to_widget_no_use(self, list_articles):
        self.ui.list_wodget_no_use_articles.clear()
        self.ui.list_wodget_use_articles.clear()
        for article in list_articles:
            item = QListWidgetItem(str(article[1]))  
            item_widget = QWidget()
            separator_name = QLabel()
            name_article = QLabel(str(article[1]))
            add_push_button = QPushButton("+")
            add_push_button.setObjectName(str(article[0]))
            add_push_button.clicked.connect(lambda checked, item=item: self.add_article_push_button(item))
            item_layout = QHBoxLayout()
            item_layout.addWidget(name_article)
            item_layout.addWidget(separator_name)
            item_layout.addWidget(add_push_button)
            item_widget.setLayout(item_layout)
            item.setSizeHint(item_widget.sizeHint())
            self.ui.list_wodget_no_use_articles.addItem(item)
            self.ui.list_wodget_no_use_articles.setItemWidget(item, item_widget)

    def add_article_push_button(self, item):
        row = self.ui.list_wodget_no_use_articles.row(item)
        if row != -1:
            article_name = item.text()
            self.ui.list_wodget_no_use_articles.takeItem(row)
            new_item = QListWidgetItem(article_name)  
            new_item_widget = QWidget()
            remove_push_button = QPushButton("-")
            remove_push_button.setObjectName(article_name)  
            separator_name = QLabel()
            remove_push_button.clicked.connect(lambda checked, item=new_item: self.remove_article_push_button(item, article_name))
            item_layout = QHBoxLayout()
            item_layout.addWidget(QLabel(article_name))  
            item_layout.addWidget(separator_name)
            item_layout.addWidget(remove_push_button)
            new_item_widget.setLayout(item_layout)
            new_item.setSizeHint(new_item_widget.sizeHint())
            self.ui.list_wodget_use_articles.addItem(new_item)
            self.ui.list_wodget_use_articles.setItemWidget(new_item, new_item_widget)

    def remove_article_push_button(self, item, article_name):
        row = self.ui.list_wodget_use_articles.row(item)
        if row != -1:
            self.ui.list_wodget_use_articles.takeItem(row)
            return_item = QListWidgetItem(article_name)  
            return_widget = QWidget()
            add_push_button = QPushButton("+")
            add_push_button.setObjectName(article_name)  
            separator_name = QLabel()
            add_push_button.clicked.connect(lambda checked, item=return_item: self.add_article_push_button(return_item))
            return_layout = QHBoxLayout()
            return_layout.addWidget(QLabel(article_name)) 
            return_layout.addWidget(separator_name)
            return_layout.addWidget(add_push_button)
            return_widget.setLayout(return_layout)
            return_item.setSizeHint(return_widget.sizeHint())
            self.ui.list_wodget_no_use_articles.addItem(return_item)
            self.ui.list_wodget_no_use_articles.setItemWidget(return_item, return_widget)

    def add_journals_to_db(self, name, data, number):
        onnection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                cursor = connection.cursor()
                query = """INSERT INTO Journals (name, date_create, number_journal) VALUES (%s, %s, %s);"""
                cursor.execute(query, (name, data, number,))
                query = f"""INSERT INTO Journals (name, date_create, number_journal) VALUES ({name}, {data}, {number});"""
                self.write_sql(query)
        except Error as e:
            QMessageBox.warning(self, "Ошибка", "Название журнала должно быть уникальным")
            return True
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
    
    def get_all_journals_from_db(self, param=None, value=None):
        onnection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                cursor = connection.cursor()
                if param and value:
                    query = f"""SELECT * FROM Journals WHERE {param} = %s;"""
                    cursor.execute(query, (value,))
                    query = f"""SELECT * FROM Journals WHERE {param} = {value};"""
                    self.write_sql(query)
                else:
                    cursor.execute("""SELECT * FROM Journals;""")
                    query = f"""SELECT * FROM Journals;"""
                    self.write_sql(query)
                items = cursor.fetchall()
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        return items

    def loading_journasl_to_list_widget_page_journals(self, list_journals):
        self.ui.list_widget_journals.clear()
        for journal in list_journals:
            item = QListWidgetItem()
            item_widget = QWidget()
            name_journal = QLabel(str(journal[1]))
            separator_name = QLabel()
            date_journal = QLabel(str(journal[2]))
            separator_date = QLabel()
            number_journal = QLabel(str(journal[3]))
            separator_button = QLabel()
            delete_push_button = QPushButton("Удалить")
            view_push_button = QPushButton("Смотреть")
            delete_push_button.setObjectName(str(journal[0]))  # Именем кнопки для удаления будет являться id журнала
            view_push_button.setObjectName(str(journal[0]))  # Именем кнопки для просмотра будет являться id журнала
            delete_push_button.clicked.connect(self.delete_journal_push_button)
            view_push_button.clicked.connect(self.view_journal_push_button)
            view_push_button.setStyleSheet("""
                QPushButton {
                    background-color: #d90429; 
                    color: white; 
                    border-radius: 5px; 
                    font-size: 17px; 
                    padding: 4px 7px;
                }
                QPushButton:pressed {
                    background-color: #ff5a5f;  /* Цвет при нажатии */
                }
            """)
            delete_push_button.setStyleSheet("""
                QPushButton {
                    background-color: #d90429; 
                    color: white; 
                    border-radius: 5px; 
                    font-size: 17px; 
                    padding: 4px 7px;
                }
                QPushButton:pressed {
                    background-color: #ff5a5f;  /* Цвет при нажатии */
                }
            """)
            item_layout = QHBoxLayout()
            item_layout.addWidget(name_journal)
            item_layout.addWidget(separator_name)
            item_layout.addWidget(date_journal)
            item_layout.addWidget(separator_date)
            item_layout.addWidget(number_journal)
            item_layout.addWidget(separator_button)
            item_layout.addWidget(delete_push_button)
            item_layout.addWidget(view_push_button)
            item_widget.setLayout(item_layout)
            item_widget.setStyleSheet("background-color: #81c3d7; font-size: 17px; border-radius: 10px; margin-bottom: 9px; margin-top: 9px; margin-right: 5px; margin-left: 5px; ")
            item.setSizeHint(item_widget.sizeHint())
            self.ui.list_widget_journals.addItem(item)
            self.ui.list_widget_journals.setItemWidget(item, item_widget)
    
    def connect_article_with_journal(self, id_journal, list_id_articles):

        for id_article in list_id_articles:
            onnection = None
            cursor = None
            try:
                connection = connect(
                    host="sql7.freesqldatabase.com",
                    user="sql7753735",
                    password="8a8kJ57nI8",
                    database="sql7753735" # Имя базы данных
                )
                if connection.is_connected():
                    cursor = connection.cursor()
                    query = """INSERT INTO Articles_Journals (ID_Journal, ID_Article) VALUES (%s, %s);"""
                    cursor.execute(query, (id_journal, id_article,))
                    query = f"""INSERT INTO Articles_Journals (ID_Journal, ID_Article) VALUES ({id_journal}, {id_article});"""
                    self.write_sql(query)
            except Error as e:
                print(f"Ошибка подключения: {e}")
            finally:
                if cursor:
                    cursor.close()
                if connection:
                    connection.commit()
                    connection.close()

    def get_id_journal(self, name, date_create, number_journal):
        onnection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                cursor = connection.cursor()
                query = """SELECT ID_journal FROM Journals WHERE name = %s AND date_create = %s AND  number_journal = %s;"""
                cursor.execute(query, (name, date_create, number_journal,))
                query = f"""SELECT ID_journal FROM Journals WHERE name = {name} AND date_create = {date_create} AND  number_journal = {number_journal};"""
                self.write_sql(query)
                id_journal = cursor.fetchone()
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        return id_journal

    def get_id_articles(self, list_name_articles):
        list_id_articles = []
        for name_article in list_name_articles:
            onnection = None
            cursor = None
            try:
                connection = connect(
                    host="sql7.freesqldatabase.com",
                    user="sql7753735",
                    password="8a8kJ57nI8",
                    database="sql7753735" # Имя базы данных
                )
                if connection.is_connected():
                    cursor = connection.cursor()
                    query = """SELECT ID_Article FROM Articles WHERE name = %s;"""
                    cursor.execute(query, (name_article,))
                    query = f"""SELECT ID_Article FROM Articles WHERE name = {name_article}""" 
                    self.write_sql(query)
                    id_article = cursor.fetchone()
                    id_article = id_article[0] 
                    list_id_articles.append(id_article)

            except Error as e:
                print(f"Ошибка подключения: {e}")
            finally:
                if cursor:
                    cursor.close()
                if connection:
                    connection.close()
        return list_id_articles

    def change_status_articles(self, list_id_articles):
        for id_article in list_id_articles:
            onnection = None
            cursor = None
            try:
                connection = connect(
                    host="sql7.freesqldatabase.com",
                    user="sql7753735",
                    password="8a8kJ57nI8",
                    database="sql7753735" # Имя базы данных
                )
                if connection.is_connected():
                    cursor = connection.cursor()
                    query = """UPDATE Articles SET isUse = 1 WHERE ID_Article = %s"""
                    cursor.execute(query, (id_article,))
                    query = f"""UPDATE Articles SET isUse = 1 WHERE ID_Article = {id_article}"""
                    self.write_sql(query)
            except Error as e:
                print(f"Ошибка подключения: {e}")
            finally:
                if cursor:
                    cursor.close()
                if connection:
                    connection.commit()
                    connection.close()

    def search_journals_with_param(self):
        param = self.ui.combo_box_journals.currentText()
        value = self.ui.line_edit_value_param_journals.text()
        if param == "Название":
            current_param = "name"
        elif param == "Дата":
            if not self.validate_date(value):
                QMessageBox.warning(self, "Ошибка", "Введите корректную дату в формате ГГГГ-ММ-ДД или РЕАЛЬНУЮ дату")
                return
            current_param = "date_create" 
        elif param == "Номер":
            current_param = "number_journal"
        elif param == "Сброс параметров":
            current_param = None
            value = None
        list_journals = self.get_all_journals_from_db(current_param, value)
        self.loading_journasl_to_list_widget_page_journals(list_journals)

    def delete_journal_push_button(self):
        sender = self.sender()
        push_button = self.findChild(QPushButton, sender.objectName())
        onnection = None
        cursor = None
        try:
            db_config = self.read_db_config('connection_db.txt')
            connection = connect(
                host=db_config["host"],
                user=db_config["user"],
                password=db_config["password"],
                database=db_config["database"]
            )
            if connection.is_connected():
                cursor = connection.cursor()
                query = """SELECT ID_Article FROM Articles_Journals WHERE ID_Journal = %s"""
                cursor.execute(query, (sender.objectName(),))
                id_articles = cursor.fetchall()
                query = f"""SELECT ID_Article FROM Articles_Journals WHERE ID_Journal = {sender.objectName()}"""
                self.write_sql(query)
                list_id_articles = []
                for el in id_articles:
                    list_id_articles.append(el[0])
                query = """DELETE FROM Journals WHERE ID_journal = %s"""
                cursor.execute(query, (sender.objectName(),))
                query = f"""DELETE FROM Journals WHERE ID_journal = {sender.objectName()}"""
                self.write_sql(query)
                for el in list_id_articles:
                    query = """UPDATE Articles SET isUse = 0 WHERE ID_Article = %s"""
                    cursor.execute(query, (el,))
                query = f"""UPDATE Articles SET isUse = 0 WHERE ID_Article = {el}"""
                self.write_sql(query)
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        all_journals = self.get_all_journals_from_db()
        self.loading_journasl_to_list_widget_page_journals(all_journals)
        list_of_articles_no_use = self.loading_all_articles_with_not_use()
        self.insert_articles_to_widget_no_use(list_of_articles_no_use)
        self.ui.list_wodget_use_articles.clear()

    def get_all_info_about_journal_wtih_param(self, id_journal):
        onnection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                cursor = connection.cursor()
                query = f"""
                        SELECT
	                        Authors.ID_Author,
	                        Authors.first_name,
	                        Authors.last_name,
	                        Articles.ID_Article,
	                        Articles.name,
	                        Articles.date_create,
	                        Articles.science,
	                        Articles.text_article,
                            Journals.name
                        FROM
	                        Articles
                        INNER JOIN
	                        Authors_Articles ON Articles.ID_Article = Authors_Articles.ID_Article
                        INNER JOIN
	                        Authors ON Authors_Articles.ID_Author = Authors.ID_Author
                        INNER JOIN
	                        Articles_Journals ON Articles.ID_Article = Articles_Journals.ID_Article
                        INNER JOIN
	                        Journals ON Journals.ID_Journal = Articles_Journals.ID_Journal
                        WHERE
	                        Journals.ID_Journal = %s; 
                        """
                # cursor.execute(query)
                # self.write_sql(query)
                cursor.execute(query, (id_journal,))
                items = cursor.fetchall()
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        return items
   
    def view_journal_push_button(self):
        sender = self.sender()
        push_button = self.findChild(QPushButton, sender.objectName())
        all_info_about_articles_of_journal_with_id = self.get_all_info_about_journal_wtih_param(sender.objectName())

        if not all_info_about_articles_of_journal_with_id:
            QMessageBox.warning(self, "Сообщение", "У данного журнала нет статей")
            return
        document = ap.Document()
        page = document.pages.add()
        header_text = ap.text.TextFragment(f"Выпуск \"{all_info_about_articles_of_journal_with_id[0][8]}\"")
        page.paragraphs.add(header_text)
        sep_text = ap.text.TextFragment(" ")
        page.paragraphs.add(sep_text)
        text = ap.text.TextFragment("Статьи, которые входят в журнал:")
        page.paragraphs.add(text)
        count = 1
        for el in all_info_about_articles_of_journal_with_id:
            text = ap.text.TextFragment(f"   {count} : " + str(el[1]) + " " + str(el[2]) + " - \"" + str(el[4]) + "\"")
            page.paragraphs.add(text)
            count += 1
        page.paragraphs.add(sep_text)
        text = ap.text.TextFragment("Статьи")
        page.paragraphs.add(text)
        page.paragraphs.add(sep_text)
        page.paragraphs.add(sep_text)
        page.paragraphs.add(sep_text)
        for el in all_info_about_articles_of_journal_with_id:
            name = ap.text.TextFragment(f"Статья - {el[4]}")
            page.paragraphs.add(name)
            author = ap.text.TextFragment(f"Автор - {str(el[1]) + " " + str(el[2])}")
            page.paragraphs.add(author)
            date = ap.text.TextFragment(f"Дата - {str(el[5])}")
            page.paragraphs.add(date)
            science = ap.text.TextFragment(f"Наука - {str(el[6])}")
            page.paragraphs.add(science)
            text = ap.text.TextFragment(f"{el[7]}")
            page.paragraphs.add(text)
            page.paragraphs.add(sep_text)
            page.paragraphs.add(sep_text)
            page.paragraphs.add(sep_text)
        project_path = os.path.dirname(os.path.abspath(sys.modules['__main__'].__file__))
        file_name = f'article_{push_button.objectName()}.pdf'  
        file_path = os.path.join(project_path, file_name)
        document.save(file_path)
        os.startfile(file_path)

    def create_db(self):
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                cursor = connection.cursor()
                query = """
                    CREATE TABLE IF NOT EXISTS Authors (
                        ID_Author INT NOT NULL AUTO_INCREMENT PRIMARY KEY,
                        first_name VARCHAR(150) NOT NULL,
                        last_name VARCHAR(150) NOT NULL,
                        middle_name VARCHAR(150) NOT NULL,
                        info TEXT NOT NULL,
                        count_articles INT NOT NULL DEFAULT 0
                    );
                """
                cursor.execute(query)
                self.write_sql(query)
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS Articles (
                        ID_Article INT NOT NULL AUTO_INCREMENT PRIMARY KEY,
                        name VARCHAR(255) NOT NULL UNIQUE,
                        date_create DATE NOT NULL,
                        science VARCHAR(255) NOT NULL,
                        text_article TEXT NOT NULL,
                        isUse BOOLEAN NOT NULL DEFAULT FALSE
                    );
                """)
                cursor.execute(query)
                self.write_sql(query)
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS Authors_Articles (
                        ID INT NOT NULL AUTO_INCREMENT PRIMARY KEY,
                        ID_Author INT,
                        ID_Article INT UNIQUE,
                        FOREIGN KEY (ID_Author) REFERENCES Authors(ID_Author) ON DELETE CASCADE,
                        FOREIGN KEY (ID_Article) REFERENCES Articles(ID_Article) ON DELETE CASCADE
                    );
                """)
                cursor.execute(query)
                self.write_sql(query)
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS Journals (
                        ID_Journal INT NOT NULL AUTO_INCREMENT PRIMARY KEY,
                        name VARCHAR(255) NOT NULL UNIQUE,
                        date_create DATE NOT NULL,
                        number_journal INT NOT NULL
                    );
                """)
                cursor.execute(query)
                self.write_sql(query)
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS Articles_Journals (
                        ID INT NOT NULL AUTO_INCREMENT PRIMARY KEY,
                        ID_Journal INT,
                        ID_Article INT UNIQUE,
                        FOREIGN KEY (ID_Journal) REFERENCES Journals(ID_Journal) ON DELETE CASCADE,
                        FOREIGN KEY (ID_Article) REFERENCES Articles(ID_Article) ON DELETE CASCADE
                    );
                """)
                cursor.execute(query)
                self.write_sql(query)
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()

    def write_sql(self, query):
        with open('data.sql', 'a') as file:
            formatted_query = "\n".join(line.strip() for line in query.strip().splitlines())
            file.write("\n")
            file.write("\n")
            file.write(formatted_query)

class EditAuthorWidget(QWidget, Ui_Form):
    def __init__(self, main_window, id_author):
        super().__init__()
        self.ui = Ui_Form()  
        self.ui.setupUi(self)  
        self.main_window = main_window
        self.id_author = id_author
        #///////////////////////////////////////////////
        self.setWindowTitle("Редактирование автора")
        self.ui.putton_edit_author_save.clicked.connect(self.save_author_info)
        info = self.get_info_about_author()
        self.ui.edit_author_first_name.setText(str(info[1]))
        self.ui.edit_author_last_name.setText(str(info[2]))
        self.ui.edit_author_middle_name.setText(str(info[3]))
        self.ui.edit_author_info.setPlainText(str(info[4]))

    def save_author_info(self):
        self.change_info_abot_author()
        list_of_authors = self.main_window.get_list_of_authors()
        self.main_window.loading_authors_from_the_database_to_page(list_of_authors)
        self.close()
    
    def read_db_config(self, filename):
        config = {}
        with open(filename, 'r') as file:
            for line in file:
                key, value = line.strip().split('=')
                config[key] = value
        return config
    
    def write_sql(self, query):
        with open('data.sql', 'a') as file:
            formatted_query = "\n".join(line.strip() for line in query.strip().splitlines())
            file.write("\n")
            file.write("\n")
            file.write(formatted_query)

    def get_info_about_author(self):
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()
                query = f"SELECT * FROM Authors WHERE ID_Author = %s;"
                cursor.execute(query, (self.id_author,))
                query = f"""SELECT * FROM Authors WHERE ID_Author = {self.id_author};"""
                self.write_sql(query)
                authors = cursor.fetchone()
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()
        return authors

    def change_info_abot_author(self):
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()
                new_first_name = str(self.ui.edit_author_first_name.text())
                new_last_name = str(self.ui.edit_author_last_name.text())
                new_middle_name = str(self.ui.edit_author_middle_name.text())
                new_info = str(self.ui.edit_author_info.toPlainText())
                query = """
                    UPDATE Authors 
                    SET first_name = %s, last_name = %s, middle_name = %s, info = %s 
                    WHERE ID_Author = %s
                """
                data = (new_first_name, new_last_name, new_middle_name, new_info, self.id_author)
                cursor.execute(query, data)
                connection.commit()
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()

class EditArticleWidget(QWidget, Ui_Form_Article):
    def __init__(self, main_window, id_article):
        super().__init__()
        self.ui = Ui_Form_Article()  
        self.ui.setupUi(self)  
        self.ui.new_date_line_edit.setInputMask('0000-00-00;_')
        self.main_window = main_window
        self.id_article = id_article
        #///////////////////////////////////////////////
        self.setWindowTitle("Редактирование статьи")
        self.ui.push_button_save_changes_article_info.clicked.connect(self.save_changes)
        info_about_article = self.get_info_about_article_with_id()
        self.id_author_article = info_about_article[0][0]
        all_authors = self.get_all_author()
        current_article = info_about_article[0]
        self.ui.new_name_line_edit.setText(str(current_article[5]))
        self.ui.new_date_line_edit.setText(str(current_article[6]))
        self.ui.new_science_line_edit.setText(str(current_article[7]))
        self.ui.new_text_edit.setText(str(current_article[8]))
        self.ui.new_author_combo_box.addItem(str(current_article[0]) + " " + str(current_article[1]) + " " + str(current_article[2]) + " " + str(current_article[3]))
        for el in all_authors:
            self.ui.new_author_combo_box.addItem(str(el[0]) + " " + str(el[1]) + " " + str(el[2]) + " " + str(el[3]))

    def change_info_about_article(self):
        new_name = self.ui.new_name_line_edit.text()
        new_data = self.ui.new_date_line_edit.text()
        new_science = self.ui.new_science_line_edit.text()
        new_text = self.ui.new_text_edit.toPlainText()
        new_name_author = self.ui.new_author_combo_box.currentText()
        new_name_author = new_name_author.split()
        new_id_author = new_name_author[0]
        new_name_author_first_name = new_name_author[1]
        new_name_author_last_name = new_name_author[2]
        new_name_author_middle_name = new_name_author[3]
        if not self.validate_date(new_data):
            QMessageBox.warning(self, "Ошибка", "Введите корректную дату в формате ГГГГ-ММ-ДД или РЕАЛЬНУЮ дату")
            self.ui.new_date_line_edit.clear()
            return
        elif not new_name:
            QMessageBox.warning(self, "Ошибка", "Введите название статьи")
            return
        elif not new_science:
            QMessageBox.warning(self, "Ошибка", "Введите название науки")
            return
        elif not new_text:
            QMessageBox.warning(self, "Ошибка", "Введите название текст")
            return
        elif len(new_name) > 255:
            QMessageBox.warning(self, "Ошибка", "Введите название статьи меньше")
            return
        elif len(new_science) > 255:
            QMessageBox.warning(self, "Ошибка", "Введите название статьи меньше")
            return
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                cursor = connection.cursor()
                query = f"UPDATE Articles SET name = %s, date_create = %s, science = %s, text_article = %s WHERE ID_Article = %s;"
                cursor.execute(query, (new_name, new_data, new_science, new_text, self.id_article,))
                query = f"""UPDATE Articles SET name = {new_name}, date_create = {new_data}, science = {new_science}, text_article = {new_text} WHERE ID_Article = {self.id_article};"""
                self.write_sql(query)
        except Error as e:
            QMessageBox.warning(self, "Ошибка", "Название статьи должно быт уникальным")
            print(f"Ошибка подключения: {e}")
            return 
        finally:
            if cursor:
                connection.commit()
                cursor.close()
            if connection:
                connection.commit()
                connection.close()

        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()
                query = f"SELECT ID_Author FROM Authors WHERE ID_Author = %s AND first_name = %s AND last_name = %s AND middle_name = %s;"
                cursor.execute(query, (new_id_author, new_name_author_first_name, new_name_author_last_name, new_name_author_middle_name,))
                id_new_author = cursor.fetchall()
                query = f"""SELECT ID_Author FROM Authors WHERE ID_Author = {new_id_author} AND first_name = {new_name_author_first_name} AND last_name = {new_name_author_last_name} AND middle_name = {new_name_author_middle_name};"""
                self.write_sql(query)
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()
        new_id_author_param = id_new_author[0][0]

        self.change_connect_author_with_article(self.id_article, self.id_author_article, new_id_author_param)

    def change_connect_author_with_article(self, id_article, old_id_author, new_id_author):
        connection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()
                query = f"DELETE FROM Authors_Articles WHERE ID_Article = %s;"
                cursor.execute(query, (id_article,))
                connection.commit()
                query = f"INSERT INTO Authors_Articles (ID_Author, ID_Article) VALUES (%s, %s);"
                cursor.execute(query, (new_id_author, id_article,))
                connection.commit()
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()     

    def save_changes(self):
        self.change_info_about_article()
        list_articles = self.main_window.get_all_articles_from_db()
        self.main_window.load_articles_from_db_to_page_articles(list_articles)
        self.close()

    def get_info_about_article_with_id(self):
        onnection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()
                cursor.execute("""
                SELECT
                    Authors.ID_Author,
                    Authors.first_name,
                    Authors.last_name,
                    Authors.middle_name,
                    Articles.ID_Article,
                    Articles.name,
                    Articles.date_create,
                    Articles.science,
                    Articles.text_article
                FROM
                    Articles
                INNER JOIN
                    Authors_Articles ON Articles.ID_Article = Authors_Articles.ID_Article
                INNER JOIN
                    Authors ON Authors_Articles.ID_Author = Authors.ID_Author
                WHERE
                    Articles.ID_Article = %s;
                """, (self.id_article,))
                items = cursor.fetchall()
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        return items

    def get_all_author(self):
        onnection = None
        cursor = None
        try:
            connection = connect(
                host="sql7.freesqldatabase.com",
                user="sql7753735",
                password="8a8kJ57nI8",
                database="sql7753735" # Имя базы данных
            )
            if connection.is_connected():
                print("Успешное подключение")
                cursor = connection.cursor()
                cursor.execute("""SELECT ID_Author, first_name, last_name, middle_name FROM Authors WHERE ID_Author <> %s;""", (self.id_author_article,))
                items = cursor.fetchall()
        except Error as e:
            print(f"Ошибка подключения: {e}")
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.commit()
                connection.close()
        
        return items
    def validate_date(self, date_string):
        if len(date_string) != 10:
            return False
        year = date_string[:4]
        month = date_string[5:7]
        day = date_string[8:]
        if int(year) >= 1000:
            if int(month) >= 1 and int(month) <= 12:
                if int(day) >= 1 and int(day) <= 31:
                    if month in ['04', '06', '09', '11'] and int(day) == 31:
                        return False  
                    if month == '02':
                        if (int(year) % 4 == 0 and int(year) % 100 != 0) or (int(year) % 400 == 0):
                            if int(day) > 29:
                                return False
                        else:
                            if int(day) > 28:
                                return False
                    return True
        return False
    
    def read_db_config(self, filename):
        config = {}
        with open(filename, 'r') as file:
            for line in file:
                key, value = line.strip().split('=')
                config[key] = value
        return config

    def write_sql(self, query):
        with open('data.sql', 'a') as file:
            formatted_query = "\n".join(line.strip() for line in query.strip().splitlines())
            file.write("\n")
            file.write("\n")
            file.write(formatted_query)

if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    window = MainWindow()
    # window.showFullScreen()
    
    # window.setFixedSize(1920, 1000)
    # window.setWindowTitle('Детские кружки')
    
    window.showMaximized() 

    window.setWindowTitle("Электронный научный журнал")

    window.show()
    sys.exit(app.exec_())